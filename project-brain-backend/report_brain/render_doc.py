"""
report_brain.render_doc — universal renderer: walk a resolved document
(from assemble.resolve_document) and emit .docx. Because the frontend renders
the SAME resolved blocks, screen == file for every family (DO / PMC / Agenda /
CAPEX / WPR). Preserves each family's structure & register.
"""
from __future__ import annotations

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

_DISC_ORDER_HINT = ["Design & Engineering", "Civil", "Structural", "Mechanical",
                    "Electrical", "Piping", "Refractory", "Instrumentation",
                    "Commissioning", "Safety", "General"]


def _table_rows(rows) -> list[list]:
    if not rows:
        return []
    if isinstance(rows, dict):
        return [list(rows.values())]
    if isinstance(rows, (list, tuple)) and rows and not isinstance(rows[0], (list, tuple, dict)):
        return [list(rows)]
    out: list[list] = []
    for row in rows:
        if isinstance(row, dict):
            out.append(list(row.values()))
        elif isinstance(row, (list, tuple)):
            out.append(list(row))
        else:
            out.append([row])
    return out


def _u(run):
    run.underline = True; run.bold = True


def render_document(resolved: dict, out_path: str) -> str:
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Pt(54)

    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _u(tp.add_run(f"{resolved['title']} — {resolved.get('project','')} — {resolved.get('month','')}"))

    for blk in resolved["blocks"]:
        k = blk["kind"]
        if k == "heading":
            p = doc.add_paragraph()
            _u(p.add_run(f"{blk['roman'] + '. ' if blk.get('roman') else ''}{blk['text']}"))
        elif k == "para":
            doc.add_paragraph(blk["text"])
        elif k == "table":
            if blk.get("title"):
                tp = doc.add_paragraph(); r = tp.add_run(blk["title"]); r.bold = True; r.font.size = Pt(10)
            cols = blk["columns"]
            tbl = doc.add_table(rows=1, cols=len(cols)); tbl.style = "Table Grid"
            for i, h in enumerate(cols):
                rr = tbl.rows[0].cells[i].paragraphs[0].add_run(str(h)); rr.bold = True; rr.font.size = Pt(8.5)
            for row in _table_rows(blk.get("rows", [])):
                cells = tbl.add_row().cells
                for i, v in enumerate(row[:len(cols)]):
                    cells[i].text = "" if v is None else str(v)
            if blk.get("note"):
                np = doc.add_paragraph(); nr = np.add_run(f"Note: {blk['note']}")
                nr.italic = True; nr.font.size = Pt(8); nr.font.color.rgb = RGBColor(0x71, 0x83, 0x9A)
        elif k == "narrative":
            if blk.get("title"):
                tp = doc.add_paragraph(); r = tp.add_run(blk["title"]); r.bold = True; r.font.size = Pt(10.5)
            cur = None
            bullets = sorted(blk["bullets"],
                             key=lambda b: _DISC_ORDER_HINT.index(b["discipline"]) if b.get("discipline") in _DISC_ORDER_HINT else 99)
            if not bullets:
                doc.add_paragraph("Nil for the month.", style="List Bullet")
            for b in bullets:
                if b.get("discipline") and b["discipline"] != cur:
                    cur = b["discipline"]
                    dp = doc.add_paragraph(); dr = dp.add_run(cur); dr.bold = True; dr.font.size = Pt(10)
                para = doc.add_paragraph(style="List Bullet")
                para.add_run(b["text"].rstrip(".") + ".").font.size = Pt(10.5)
                if not b.get("grounded", True):
                    fr = para.add_run("  [unverified]"); fr.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B); fr.font.size = Pt(8)
                if b.get("draft"):
                    fr = para.add_run("  [auto-draft]"); fr.font.color.rgb = RGBColor(0xB7, 0x79, 0x1F); fr.font.size = Pt(8)

    doc.save(out_path)
    return out_path
