"""
report_brain.api — FastAPI surface for Report Studio.

  POST /report-brain/ingest        upload a source file -> atoms into the store
  POST /report-brain/quick-note    type a note directly -> atom (first-class source)
  GET  /report-brain/projects      projects with atom counts for a month
  POST /report-brain/compose       compose all sections for a project/month
  GET  /report-brain/review        get composed sections with per-bullet citations
  POST /report-brain/edit          human correction -> store + taught-facts loop
  POST /report-brain/generate      render the report family -> downloadable file
  GET  /report-brain/commitments   commitment lifecycle board for a project

The store is process-global (MemStore for dev; swap to the DB repo in prod —
same surface). Every compose is grounded; every edit is captured for learning.
"""
from __future__ import annotations

import os
import tempfile
import time
import hashlib
import json
from datetime import datetime
from typing import Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse
from pydantic import BaseModel
from sqlalchemy import text
from sqlalchemy.orm import Session

from report_brain.atoms import AliasRegistry, Atom
from report_brain.whatsapp import parse_whatsapp
from report_brain.recordnotes import parse_record_notes
from report_brain.dpr import extract_dpr
from report_brain.factstore import MemStore, ensure_schema
from report_brain.composer import compose_present_status, render_prose
from report_brain.issues_actions import (
    classify_commitments, compose_actions, compose_issues,
)
from report_brain.figures import manpower_average
from report_brain.render_pmc import render_pmc
from report_brain.docx_reference import parse_docx_reference
from app.core.database import get_db

router = APIRouter(prefix="/report-brain", tags=["report-brain"])

ensure_schema()

REG = AliasRegistry(path=os.environ.get("RB_ALIASES", "report_brain_aliases.json"))
STORE = MemStore()
_COMMITMENTS: dict[str, list] = {}          # project -> classified commitments
OUT_DIR = os.environ.get("RB_OUT", tempfile.gettempdir())
REF_DIR = os.path.join(OUT_DIR, "report_brain_refs")
os.makedirs(REF_DIR, exist_ok=True)


def _ref_title(family: str, project: str, month: str) -> str:
    return f"{family.upper()} - {project} - {month}"


def _doc_to_reference_text(resolved: dict) -> str:
    lines: list[str] = []
    for block in resolved.get("blocks", []):
        kind = block.get("kind")
        if kind == "heading":
            lines.append(block.get("text", ""))
        elif kind == "para":
            lines.append(block.get("text", ""))
        elif kind == "table":
            if block.get("title"):
                lines.append(block["title"])
            for row in block.get("rows", []):
                lines.append(" | ".join("" if value is None else str(value) for value in row))
        elif kind == "narrative":
            if block.get("title"):
                lines.append(block["title"])
            for bullet in block.get("bullets", []):
                lines.append(bullet.get("text", ""))
    return "\n".join(line for line in lines if line).strip()


def _latest_reference(db: Session, family: str, project: str):
    return db.execute(text("""
        SELECT id, family, project, month, title, source_kind, status, resolved,
               reference_text, file_name, file_path, notes, created_at, updated_at
        FROM rb_reference_reports
        WHERE family = :f AND project = :p AND status = 'approved'
        ORDER BY month DESC, created_at DESC
        LIMIT 1
    """), {"f": family, "p": project}).mappings().first()


def _mark_document_changes(current: dict, baseline: dict | None) -> tuple[dict, dict | None]:
    if not baseline:
        return current, None

    changed_blocks = 0
    baseline_blocks = baseline.get("blocks", [])
    current_blocks = current.get("blocks", [])
    used_baseline: set[int] = set()
    max_len = max(len(current_blocks), len(baseline_blocks))

    for index in range(max_len):
        if index >= len(current_blocks):
            break
        cur = current_blocks[index]
        base = _best_baseline_block(cur, baseline_blocks, used_baseline, index)
        if not base:
            cur["changed"] = True
            changed_blocks += 1
            continue

        kind = cur.get("kind")
        if kind == "para":
            cur["changed"] = (cur.get("text", "").strip() != base.get("text", "").strip())
        elif kind == "heading":
            cur["changed"] = (
                cur.get("text", "").strip() != base.get("text", "").strip()
                or cur.get("roman", "") != base.get("roman", "")
            )
        elif kind == "table":
            changed_cells: list[str] = []
            cur_rows = cur.get("rows", [])
            base_rows = base.get("rows", [])
            for ri, row in enumerate(cur_rows):
                for ci, value in enumerate(row):
                    base_value = None
                    if ri < len(base_rows) and ci < len(base_rows[ri]):
                        base_value = base_rows[ri][ci]
                    if str(value or "").strip() != str(base_value or "").strip():
                        changed_cells.append(f"{ri}:{ci}")
            cur["changed_cells"] = changed_cells
            cur["changed"] = bool(changed_cells)
        elif kind == "narrative":
            base_bullets = base.get("bullets", [])
            for bi, bullet in enumerate(cur.get("bullets", [])):
                base_text = base_bullets[bi].get("text", "") if bi < len(base_bullets) else ""
                bullet["changed"] = bullet.get("text", "").strip() != base_text.strip()
            cur["changed"] = any(b.get("changed") for b in cur.get("bullets", []))
        else:
            cur["changed"] = False

        if cur.get("changed"):
            changed_blocks += 1

    meta = {
        "title": baseline.get("title") or "",
        "month": baseline.get("month") or "",
        "source_kind": baseline.get("source_kind") or "approved",
        "changed_blocks": changed_blocks,
        "has_structured_baseline": bool(baseline.get("resolved")),
    }
    return current, meta


def _norm_text(value) -> str:
    return " ".join(str(value or "").lower().split())


def _block_signature(block: dict) -> str:
    kind = block.get("kind")
    if kind == "table":
        return _norm_text(block.get("title") or " ".join(str(c) for c in block.get("columns", [])[:4]))
    if kind == "narrative":
        return _norm_text(block.get("title") or block.get("section"))
    return _norm_text(block.get("text") or block.get("title"))


def _best_baseline_block(cur: dict, baseline_blocks: list[dict], used: set[int], fallback_index: int) -> dict | None:
    cur_kind = cur.get("kind")
    cur_sig = _block_signature(cur)
    best_index = None
    best_score = 0.0
    from difflib import SequenceMatcher

    for index, block in enumerate(baseline_blocks):
        if index in used or block.get("kind") != cur_kind:
            continue
        score = SequenceMatcher(None, cur_sig, _block_signature(block)).ratio() if cur_sig else 0.0
        if score > best_score:
            best_score = score
            best_index = index

    if best_index is not None and best_score >= 0.62:
        used.add(best_index)
        return baseline_blocks[best_index]

    if fallback_index < len(baseline_blocks):
        block = baseline_blocks[fallback_index]
        if fallback_index not in used and block.get("kind") == cur_kind:
            used.add(fallback_index)
            return block
    return None


def _month_end(month: str) -> str:
    try:
        year, mon = int(month[:4]), int(month[5:7])
        import calendar
        return f"{year:04d}-{mon:02d}-{calendar.monthrange(year, mon)[1]:02d}"
    except Exception:
        return f"{month}-30"


def _project_scheme_id(db: Session, project: str) -> int | None:
    """Resolve Report Studio project labels to the live portfolio scheme ids."""
    raw = str(project or "").strip()
    if raw.isdigit():
        return int(raw)
    compact = raw.lower().replace(" ", "").replace("-", "").replace("_", "")
    aliases = {
        "cob7": 74,
        "cob7pkg2": 74,
        "cob7batteryproper": 74,
        "cob7batteryproperpkg2": 74,
        "ts2": 10,
        "treatmentsystem2": 10,
        "bf5stove4": 9,
    }
    if compact in aliases:
        return aliases[compact]
    token = raw.replace("#", "").replace("-", " ").replace("_", " ").strip()
    if not token:
        return None
    row = db.execute(text("""
        SELECT scheme_id
        FROM scheme_master
        WHERE NOT COALESCE(is_deleted, false)
          AND scheme_name ILIKE :q
        ORDER BY scheme_id
        LIMIT 1
    """), {"q": f"%{token}%"}).first()
    return int(row[0]) if row else None


def _flatten_capex_rows(rows: list[dict]) -> list[list]:
    out: list[list] = []
    for row in rows:
        out.append([
            row.get("no", ""),
            row.get("category", ""),
            row.get("totalCost", 0),
            row.get("expenditureLastFy", 0),
            row.get("capexCurrentFy", 0),
            row.get("capexCurrentFy", 0),
            row.get("expenditureCurrentFy", 0),
            round((float(row.get("expenditureCurrentFy") or 0) / float(row.get("capexCurrentFy") or 1)) * 100, 2),
        ])
        for child in row.get("point3Rows") or []:
            out.extend(_flatten_capex_rows([child]))
        for child in row.get("childRows") or []:
            out.extend(_flatten_capex_rows([child]))
    return out


def _live_document_figures(db: Session, d: "DocIn") -> dict:
    """Populate missing document figure tables from the same live DB reports.

    Frontend-supplied figures still win, but an empty/stale context no longer
    leaves Report Studio documents blank while the portfolio reports have data.
    """
    ctx = dict(d.figures_ctx or {})
    scheme_id = _project_scheme_id(db, d.project)

    if scheme_id and (d.family in ("pmc", "wpr")) and not ctx.get("pmc_discipline"):
        try:
            from app.api.v1.mos_reports import pmc_project_detail
            detail = pmc_project_detail(scheme_id=scheme_id, month=d.month, db=db)
            ctx["pmc_discipline"] = [
                [
                    row.get("item") or row.get("activity") or "",
                    row.get("overallTarget", 0),
                    row.get("cumulativePrevious", 0),
                    row.get("achievementMonth", 0),
                ]
                for row in (detail.get("physicalProgress") or [])
            ]
            manpower_rows = (detail.get("manpower") or {}).get("rows") or []
            if manpower_rows and not ctx.get("manpower"):
                ctx["manpower"] = [
                    {
                        "category": " / ".join(
                            part for part in [
                                r.get("agency", "").replace("\n", " ").strip(),
                                r.get("manpower", "").strip(),
                                r.get("category", "").strip(),
                            ] if part
                        ) or "Manpower",
                        "average": float(r.get("value") or 0),
                        "days": int((detail.get("manpower") or {}).get("filledDays") or 0),
                    }
                    for r in manpower_rows
                ]
        except Exception:
            pass

    if d.family in ("do", "agenda", "capex") and (
        not ctx.get("capex_mos") or not ctx.get("portfolio_status") or not ctx.get("scheme_master")
    ):
        try:
            from app.api.v1.mos_reports import capex_project_detail, mos_capex_summary
            summary = mos_capex_summary(report_month=d.month, db=db)
            rows = summary.get("rows") or []
            if not ctx.get("capex_mos"):
                ctx["capex_mos"] = _flatten_capex_rows(rows)
            if not ctx.get("capex_heads"):
                ctx["capex_heads"] = [
                    [
                        row.get("category", ""),
                        row.get("expenditureCurrentFy", 0),
                        row.get("capexCurrentFy", 0),
                        row.get("totalExpenditure", 0),
                        row.get("capexCurrentFy", 0),
                    ]
                    for row in rows if row.get("section") or row.get("category") in ("Total", "Spares & Capital Repairs")
                ]
            if not ctx.get("portfolio_status"):
                ongoing = next((r for r in rows if r.get("no") == "3"), {})
                groups = ongoing.get("statusGroups") or []
                if groups:
                    ctx["portfolio_status"] = [
                        [g.get("label", ""), g.get("count", 0), g.get("cost", 0)]
                        for g in groups
                    ]
                else:
                    buckets: dict[str, dict] = {}
                    for project_row in ongoing.get("projectRows") or []:
                        label = project_row.get("derivedStatus") or "Not classified"
                        bucket = buckets.setdefault(label, {"count": 0, "cost": 0.0})
                        bucket["count"] += 1
                        bucket["cost"] += float(project_row.get("totalCost") or 0)
                    ctx["portfolio_status"] = [
                        [label, value["count"], round(value["cost"], 2)]
                        for label, value in buckets.items()
                    ]
            if not ctx.get("scheme_master"):
                detail = capex_project_detail(month=d.month, db=db)
                ctx["scheme_master"] = [
                    [
                        i + 1,
                        row.get("name", ""),
                        row.get("totalCost", 0),
                        row.get("approvalDate") or "",
                        row.get("originalCompletionDate") or "",
                        row.get("cumulativeExpenditure", 0),
                        row.get("revisedCompletionDate") or "",
                        row.get("anticipatedCompletionDate") or "",
                    ]
                    for i, row in enumerate(detail.get("highCostProjects") or [])
                ]
        except Exception:
            pass

    ctx.setdefault("as_on", _month_end(d.month))
    return ctx


# --------------------------------------------------------------- ingest
@router.post("/ingest")
async def ingest(file: UploadFile = File(...), month: str = Form(...),
                 kind: str = Form("auto"), project: str = Form("")) -> dict:
    raw = await file.read()
    name = file.filename or "upload"
    suffix = name.rsplit(".", 1)[-1].lower()
    added = {"whatsapp": 0, "status": 0, "manpower": 0, "actions": 0, "issues": 0, "commitments": 0}
    staging: list[dict] = []

    if suffix == "txt":
        atoms = parse_whatsapp(raw.decode("utf-8", "ignore"), REG, month=month)
        STORE.add_atoms(atoms, month)
        added["whatsapp"] = len(atoms)
    elif suffix in ("xlsx", "xls"):
        tmp = os.path.join(OUT_DIR, f"_ing_{int(time.time()*1000)}.{suffix}")
        open(tmp, "wb").write(raw)
        res = extract_dpr(tmp, REG, as_on=f"{month}-30")
        STORE.add_atoms(res.status_atoms, month)
        STORE.add_atoms(res.manpower_atoms, month)
        added["status"] = len(res.status_atoms)
        added["manpower"] = len(res.manpower_atoms)
        staging = [s.__dict__ for s in res.progress]  # for S-curve confirm screen
        os.remove(tmp)
    elif suffix in ("docx", "txt_rn"):
        if suffix == "docx":
            from docx import Document

            tmp = os.path.join(OUT_DIR, f"_ing_{int(time.time()*1000)}.{suffix}")
            open(tmp, "wb").write(raw)
            try:
                text = "\n".join(p.text for p in Document(tmp).paragraphs)
            finally:
                os.remove(tmp)
        else:
            text = raw.decode("utf-8", "ignore")
        notes = parse_record_notes(text, project=project or "UNKNOWN")
        STORE.add_atoms(notes, month)
        for a in notes:
            added[a.kind if a.kind in added else "actions"] = added.get(a.kind, 0) + 1
    else:
        raise HTTPException(422, f"unsupported source type: {suffix}")

    return {"file": name, "month": month, "added": added,
            "staging_rows": staging, "total_atoms": len(STORE.atoms)}


class QuickNote(BaseModel):
    project: str
    month: str
    text: str
    section: str = "status"       # status | issue | action
    discipline: str = ""
    date: str = ""


@router.post("/quick-note")
async def quick_note(n: QuickNote) -> dict:
    a = Atom(kind=n.section, date=n.date or f"{n.month}-15", project=n.project,
             section_affinity=n.section, discipline=n.discipline, text=n.text.strip(),
             source_type="manual", source_ref=f"manual:{int(time.time())}", author="user")
    STORE.add_atoms([a], n.month)
    return {"added": True, "project": n.project, "total_atoms": len(STORE.atoms)}


@router.get("/projects")
async def projects(month: str) -> dict:
    counts: dict[str, dict] = {}
    for a in STORE.atoms:
        if a["month"] != month:
            continue
        p = counts.setdefault(a["project"], {"status": 0, "issue": 0, "action": 0, "manpower": 0})
        aff = a.get("section_affinity", "status")
        p[aff] = p.get(aff, 0) + 1
    return {"month": month, "projects": counts}


# --------------------------------------------------------------- compose
class ComposeIn(BaseModel):
    project: str
    month: str
    as_of: str = ""


@router.post("/compose")
async def compose(c: ComposeIn) -> dict:
    status_atoms = STORE.atoms_for(c.project, c.month, "status")
    action_atoms = STORE.atoms_for(c.project, c.month, "action")
    issue_atoms = STORE.atoms_for(c.project, c.month, "issue")
    manpower_atoms = STORE.atoms_for(c.project, c.month, "manpower")
    commit_atoms = [a for a in STORE.atoms_for(c.project, c.month)
                    if a.get("kind") == "commitment"]

    classified = classify_commitments(commit_atoms, status_atoms,
                                      as_of=c.as_of or f"{c.month}-30")
    _COMMITMENTS[c.project] = classified

    present = compose_present_status(status_atoms)
    actions = compose_actions(action_atoms, classified)
    issues = compose_issues(issue_atoms, classified)
    manpower = manpower_average(manpower_atoms)

    STORE.put_narrative(c.project, c.month, "present_status", present)
    STORE.put_narrative(c.project, c.month, "issues", issues)
    STORE.put_narrative(c.project, c.month, "actions", actions)

    return {
        "project": c.project, "month": c.month,
        "present_status": present, "issues": issues, "actions": actions,
        "manpower": manpower,
        "grounding": {
            "present_grounded": sum(b["grounded"] for b in present),
            "present_total": len(present),
            "auto_slippage": sum(1 for b in issues if b.get("draft")),
        },
        "commitments": {
            "met": sum(1 for x in classified if x["lifecycle"] == "met"),
            "open": sum(1 for x in classified if x["lifecycle"] == "open"),
            "missed": sum(1 for x in classified if x["lifecycle"] == "missed"),
        },
    }


# --------------------------------------------------------------- edit (learn)
class EditIn(BaseModel):
    project: str
    month: str
    section_type: str
    before_text: str
    after_text: str
    kind: str = "phrasing"        # phrasing | fact | structure


@router.post("/edit")
async def edit(e: EditIn) -> dict:
    """Persist a correction and feed the taught-facts learning loop."""
    learned = None
    try:
        from app.tools.memory_tools import remember_fact
        if e.kind == "fact":
            remember_fact(subject=f"{e.project} report fact",
                          fact=f"{e.before_text} -> {e.after_text}", authority=True)
            learned = "fact"
        elif e.kind == "phrasing":
            remember_fact(subject=f"{e.project} phrasing preference",
                          fact=f"prefer '{e.after_text}' over '{e.before_text}'", authority=False)
            learned = "phrasing"
    except Exception:
        pass
    # update the stored narrative bullet in place
    key = (e.project, e.month, e.section_type)
    nar = STORE.narratives.get(key)
    if nar:
        for b in nar["body"]:
            if b.get("text") == e.before_text:
                b["text"] = e.after_text
                b["edited"] = True
    return {"saved": True, "learned": learned}


@router.get("/commitments")
async def commitments(project: str) -> dict:
    return {"project": project, "commitments": _COMMITMENTS.get(project, [])}


# --------------------------------------------------------------- generate
class GenerateIn(BaseModel):
    project: str
    project_name: str
    month: str
    month_label: str
    family: str = "pmc"           # pmc | do | agenda | capex | wpr
    progress_rows: list = []
    milestones: list = []
    officials: str = ""


@router.post("/generate")
async def generate(g: GenerateIn) -> dict:
    present = STORE.narratives.get((g.project, g.month, "present_status"), {}).get("body", [])
    issues = STORE.narratives.get((g.project, g.month, "issues"), {}).get("body", [])
    actions = STORE.narratives.get((g.project, g.month, "actions"), {}).get("body", [])
    manpower = manpower_average(STORE.atoms_for(g.project, g.month, "manpower"))

    fname = f"{g.family.upper()}_{g.project}_{g.month}.docx".replace(" ", "_")
    out_path = os.path.join(OUT_DIR, fname)
    if g.family == "pmc":
        render_pmc(out_path, project_name=g.project_name, month_label=g.month_label,
                   progress_rows=g.progress_rows, present_status=present,
                   issues=issues, actions=actions, milestones=g.milestones,
                   manpower=manpower, officials=g.officials)
    else:
        # other families share the store; renderers land in later increments
        render_pmc(out_path, project_name=g.project_name, month_label=g.month_label,
                   progress_rows=g.progress_rows, present_status=present,
                   issues=issues, actions=actions, milestones=g.milestones,
                   manpower=manpower, officials=g.officials)
    return {"file": fname, "path": out_path, "download": f"/report-brain/download/{fname}"}


@router.get("/download/{fname}")
async def download(fname: str):
    path = os.path.join(OUT_DIR, os.path.basename(fname))
    if not os.path.exists(path):
        raise HTTPException(404, "file not found")
    return FileResponse(path, filename=fname,
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# --------------------------------------------------------------- live document
from report_brain.assemble import resolve_document
from report_brain.render_doc import render_document


class DocIn(BaseModel):
    family: str = "pmc"
    project: str
    month: str
    all_projects: list = []
    figures_ctx: dict = {}
    project_names: dict = {}
    as_of: str = ""


@router.post("/document")
async def document(d: DocIn, db: Session = Depends(get_db)) -> dict:
    """Return the fully resolved, editable document for a family (screen == file)."""
    figures_ctx = _live_document_figures(db, d)
    resolved = resolve_document(d.family, d.project, d.month, STORE,
                                figures_ctx=figures_ctx, project_names=d.project_names,
                                all_projects=d.all_projects or None, as_of=d.as_of)
    baseline = _latest_reference(db, d.family, d.project)
    resolved, meta = _mark_document_changes(resolved, dict(baseline) if baseline and baseline.get("resolved") else None)
    if meta:
        resolved["reference"] = meta
    return resolved


class DocExport(BaseModel):
    resolved: dict
    filename: str = "report.docx"


@router.post("/document/export")
async def document_export(x: DocExport) -> dict:
    fname = os.path.basename(x.filename)
    path = os.path.join(OUT_DIR, fname)
    render_document(x.resolved, path)
    return {"file": fname, "download": f"/report-brain/download/{fname}"}


class ApproveIn(BaseModel):
    family: str
    project: str
    month: str
    resolved: dict
    notes: str = ""


@router.post("/document/approve")
async def document_approve(x: ApproveIn, db: Session = Depends(get_db)) -> dict:
    db.execute(text("""
        UPDATE rb_reference_reports
        SET status = 'superseded', updated_at = now()
        WHERE family = :f AND project = :p AND status = 'approved'
    """), {"f": x.family, "p": x.project})
    db.execute(text("""
        INSERT INTO rb_reference_reports
            (family, project, month, title, source_kind, status, resolved, reference_text, notes, updated_at)
        VALUES
            (:f, :p, :m, :t, 'approved', 'approved', CAST(:resolved AS jsonb), :ref, :notes, now())
    """), {
        "f": x.family,
        "p": x.project,
        "m": x.month,
        "t": _ref_title(x.family, x.project, x.month),
        "resolved": json.dumps(x.resolved),
        "ref": _doc_to_reference_text(x.resolved),
        "notes": x.notes,
    })
    db.commit()
    return {"ok": True, "family": x.family, "project": x.project, "month": x.month}


@router.get("/document/reference")
async def document_reference(family: str, project: str, db: Session = Depends(get_db)) -> dict:
    ref = _latest_reference(db, family, project)
    if not ref:
        return {"reference": None}
    return {"reference": {
        "family": ref["family"],
        "project": ref["project"],
        "month": ref["month"],
        "title": ref["title"],
        "source_kind": ref["source_kind"],
        "file_name": ref["file_name"],
        "updated_at": ref["updated_at"].isoformat() if ref.get("updated_at") else None,
        "has_structured_baseline": bool(ref.get("resolved")),
    }}


@router.post("/document/upload-reference")
async def document_upload_reference(
    family: str = Form(...),
    project: str = Form(...),
    month: str = Form(...),
    notes: str = Form(""),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
) -> dict:
    name = file.filename or "reference.docx"
    if not name.lower().endswith((".docx", ".doc")):
        raise HTTPException(status_code=400, detail="Only .docx/.doc reference files are supported.")

    raw = await file.read()
    stamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")
    file_hash = hashlib.sha1(raw).hexdigest()[:12]
    safe_name = f"{family}_{project}_{month}_{stamp}_{file_hash}_{os.path.basename(name)}".replace(" ", "_")
    path = os.path.join(REF_DIR, safe_name)
    with open(path, "wb") as fh:
        fh.write(raw)

    reference_text = ""
    resolved_reference = None
    try:
        from docx import Document
        doc = Document(path)
        reference_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    except Exception:
        reference_text = ""
    try:
        resolved_reference = parse_docx_reference(
            path,
            family=family,
            project=project,
            month=month,
            title=_ref_title(family, project, month),
        )
        parsed_text = _doc_to_reference_text(resolved_reference)
        if parsed_text:
            reference_text = parsed_text
    except Exception:
        resolved_reference = None

    db.execute(text("""
        UPDATE rb_reference_reports
        SET status = 'superseded', updated_at = now()
        WHERE family = :f AND project = :p AND status = 'approved'
    """), {"f": family, "p": project})
    db.execute(text("""
        INSERT INTO rb_reference_reports
            (family, project, month, title, source_kind, status, resolved, reference_text, file_name, file_path, notes, updated_at)
        VALUES
            (:f, :p, :m, :t, 'uploaded', 'approved', CAST(:resolved AS jsonb), :ref, :fn, :fp, :notes, now())
    """), {
        "f": family,
        "p": project,
        "m": month,
        "t": _ref_title(family, project, month),
        "resolved": json.dumps(resolved_reference) if resolved_reference else None,
        "ref": reference_text,
        "fn": name,
        "fp": path,
        "notes": notes,
    })
    db.commit()

    # Feed uploaded corporate-final text back into the report store for future composition.
    if reference_text.strip():
        notes_atoms = parse_record_notes(reference_text, project=project)
        STORE.add_atoms(notes_atoms, month)

    return {
        "ok": True,
        "family": family,
        "project": project,
        "month": month,
        "file": name,
        "structured_blocks": len(resolved_reference.get("blocks", [])) if resolved_reference else 0,
    }
