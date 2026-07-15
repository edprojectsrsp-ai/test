"""
report_brain.render_pmc — render a PMC Monthly Progress Report as .docx,
matching the RSP house format (underlined section headings, progress table,
present-status / issues / actions narrative, milestones, manpower, officials).

Reads everything from the Fact Store (facts + narratives + masters); the
composer/lifecycle already produced the narrative bullets with grounding.
Ungrounded or auto-draft bullets are marked so the reviewer sees them.
"""
from __future__ import annotations

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _u(run):
    run.underline = True
    run.bold = True


def _heading(doc, roman, text):
    p = doc.add_paragraph()
    r = p.add_run(f"{roman}. {text}")
    _u(r); r.font.size = Pt(11)
    return p


def _disc_bullets(doc, bullets):
    cur = None
    for b in bullets:
        if b["discipline"] != cur:
            cur = b["discipline"]
            p = doc.add_paragraph()
            r = p.add_run(cur); r.bold = True; r.font.size = Pt(10.5)
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(b["text"] + ".")
        run.font.size = Pt(10.5)
        if not b.get("grounded", True):
            f = para.add_run("  [unverified]"); f.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B); f.font.size = Pt(8)
        if b.get("draft"):
            f = para.add_run("  [auto-draft — review]"); f.font.color.rgb = RGBColor(0xB7, 0x79, 0x1F); f.font.size = Pt(8)


def render_pmc(out_path: str, *, project_name: str, month_label: str,
               progress_rows: list[dict], present_status: list[dict],
               issues: list[dict], actions: list[dict],
               milestones: list[dict], manpower: list[dict],
               officials: str = "") -> str:
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Pt(54)

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(month_label); _u(r)
    t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run(project_name); _u(r2)

    _heading(doc, "I", "Progress of the Project")
    if progress_rows:
        tbl = doc.add_table(rows=1, cols=4); tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["Brief Description of Progress (Main package)",
                               "Overall % Target till the month", "Cumulative % completion till the month",
                               "% achievement for the month"]):
            rr = hdr[i].paragraphs[0].add_run(h); rr.bold = True; rr.font.size = Pt(9)
        for row in progress_rows:
            cells = tbl.add_row().cells
            cells[0].text = str(row.get("discipline", ""))
            cells[1].text = f"{row.get('target_till_month', '')}%"
            cells[2].text = f"{row.get('cumulative_pct', '')}%"
            cells[3].text = f"{row.get('month_pct', '')}%"

    _heading(doc, "II", "Present Status of the Project")
    _disc_bullets(doc, present_status)

    _heading(doc, "III", "Reasons / Issues / Constraints")
    if issues:
        _disc_bullets(doc, issues)
    else:
        doc.add_paragraph("No major issues reported for the month.", style="List Bullet")

    _heading(doc, "IV", "Action Taken")
    _disc_bullets(doc, actions)

    _heading(doc, "V", "Milestones reported on OCMS portal of MoSPI")
    if milestones:
        mt = doc.add_table(rows=1, cols=4); mt.style = "Table Grid"
        for i, h in enumerate(["Milestone", "Orig. Completion", "Anticipated Completion", "Reasons"]):
            rr = mt.rows[0].cells[i].paragraphs[0].add_run(h); rr.bold = True; rr.font.size = Pt(9)
        for m in milestones:
            c = mt.add_row().cells
            c[0].text = m.get("name", ""); c[1].text = m.get("orig", "")
            c[2].text = m.get("anticipated", ""); c[3].text = m.get("reason", "")

    _heading(doc, "VI", "Manpower Engaged")
    if manpower:
        mp = doc.add_table(rows=1, cols=3); mp.style = "Table Grid"
        for i, h in enumerate(["Agency / Category", "Average Engaged", "Reporting Days"]):
            rr = mp.rows[0].cells[i].paragraphs[0].add_run(h); rr.bold = True; rr.font.size = Pt(9)
        for m in manpower:
            c = mp.add_row().cells
            c[0].text = str(m.get("category", "")); c[1].text = str(int(m.get("average", 0)))
            c[2].text = str(m.get("days", ""))

    if officials:
        _heading(doc, "VII", "Key Officials for the Project")
        doc.add_paragraph(officials)

    doc.save(out_path)
    return out_path
