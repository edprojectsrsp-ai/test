"""
Report Documents — upload, view, edit-in-place, export (Word/PDF).

Stores .docx files as HTML in `record_notes` (note_type='report_doc').
On upload: mammoth converts .docx → HTML → stored as `body`.
On edit:   PUT updates the `body` HTML.
On export: htmldocx converts HTML → .docx; reportlab for PDF.

No schema change needed — uses the existing record_notes table.

Mount at:  /api/v1/report-docs
Depends:   mammoth, python-docx, htmldocx, reportlab (all pip-installable)
"""

import csv
import io
import json
from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Body, Query
from fastapi.responses import StreamingResponse
from sqlalchemy import text
from sqlalchemy.orm import Session
from app.core.database import get_db
from app.security.auth import require_user

# Sprint 0 — report document store requires auth.
router = APIRouter(tags=["Report Documents"], dependencies=[Depends(require_user)])


def _normalize_fy(value: str) -> str:
    raw = str(value or "").replace("FY", "").strip()
    if len(raw) == 9 and raw[4] == "-":
        return f"{raw[:4]}-{raw[-2:]}"
    return raw


def _capex_project_rows(db: Session, fy: str) -> list[dict]:
    rows = db.execute(text("""
        WITH selected_plan AS (
            SELECT id
            FROM capex_plan_header
            WHERE fy_year = :fy AND plan_status <> 'Archived'
            ORDER BY is_effective DESC, id DESC
            LIMIT 1
        ), leaf_rows AS (
            SELECT r.id, r.scheme_id
            FROM capex_plan_rows r
            JOIN selected_plan p ON p.id = r.plan_id
            WHERE r.row_level IN ('Item', 'Package')
              AND r.scheme_id IS NOT NULL
              AND NOT EXISTS (
                  SELECT 1 FROM capex_plan_rows child
                  WHERE child.parent_row_id = r.id
              )
        )
        SELECT s.scheme_id, s.scheme_name, s.scheme_type,
               COALESCE(SUM(v.gross_cost), 0) AS gross_cost,
               COALESCE(SUM(v.cumulative_exp_till_last_fy), 0) AS cum_last_fy,
               COALESCE(SUM(v.be_fy), 0) AS be_fy,
               COALESCE(SUM(v.re_fy), 0) AS re_fy,
               COALESCE(SUM((
                   SELECT COALESCE(SUM(a.amount), 0)
                   FROM capex_actuals a
                   WHERE a.plan_row_id = lr.id AND a.fy_year = :fy
               )), 0) AS actual_fy
        FROM leaf_rows lr
        JOIN scheme_master s ON s.scheme_id = lr.scheme_id AND NOT s.is_deleted
        JOIN capex_plan_values v ON v.plan_row_id = lr.id
        GROUP BY s.scheme_id, s.scheme_name, s.scheme_type
        ORDER BY s.scheme_type, s.scheme_name
    """), {"fy": fy}).mappings().all()
    return [dict(row) for row in rows]


def _capex_report_rows(db: Session, fy: str) -> list[dict]:
    rows = _capex_project_rows(db, fy)
    extras = db.execute(text("""
        WITH selected_plan AS (
            SELECT id
            FROM capex_plan_header
            WHERE fy_year = :fy AND plan_status <> 'Archived'
            ORDER BY is_effective DESC, id DESC
            LIMIT 1
        )
        SELECT r.row_name AS scheme_name, r.display_order,
               COALESCE(v.gross_cost, 0) AS gross_cost,
               COALESCE(v.cumulative_exp_till_last_fy, 0) AS cum_last_fy,
               COALESCE(v.be_fy, 0) AS be_fy,
               COALESCE(v.re_fy, 0) AS re_fy,
               COALESCE((
                   SELECT SUM(a.amount) FROM capex_actuals a
                   WHERE a.plan_row_id = r.id AND a.fy_year = :fy
               ), 0) AS actual_fy
        FROM capex_plan_rows r
        JOIN selected_plan p ON p.id = r.plan_id
        JOIN capex_plan_values v ON v.plan_row_id = r.id
        WHERE (
            r.indent_level = 0
            AND (r.row_name ILIKE '%MEP%' OR r.row_name ILIKE '%Capital Repairs%')
        ) OR (
            r.scheme_id IS NULL
            AND r.row_level IN ('Item', 'Package')
            AND NOT EXISTS (SELECT 1 FROM capex_plan_rows child WHERE child.parent_row_id = r.id)
        )
        ORDER BY r.display_order
    """), {"fy": fy}).mappings().all()
    return [dict(row) for row in extras] + rows


def _physical_progress_by_scheme(db: Session) -> dict[int, dict]:
    from app.api.v1.s_curve import get_scheme_package_curves

    scheme_ids = db.execute(text("""
        SELECT DISTINCT p.scheme_id
        FROM progress_plans pp
        JOIN packages p ON p.package_id = pp.package_id
        WHERE pp.is_current AND pp.is_locked AND NOT pp.is_deleted AND NOT p.is_deleted
    """)).scalars().all()
    result = {}
    for scheme_id in scheme_ids:
        packages = get_scheme_package_curves(int(scheme_id), db)
        if not packages:
            continue
        all_months = sorted({point["month_date"] for package in packages for point in package["points"]})
        actual_months = [
            point["month_date"] for package in packages for point in package["points"]
            if point["cumulative_actual_pct"] is not None
        ]
        if not all_months or not actual_months:
            continue
        cutoff = max(actual_months)
        total_weight = sum(float(package["weight"] or 0) for package in packages) or 1.0

        def value_at(points, key):
            value = 0.0
            for point in points:
                if point["month_date"] > cutoff:
                    break
                candidate = point.get(key)
                if candidate is not None:
                    value = float(candidate)
            return value

        planned = sum(
            float(package["weight"] or 0) * value_at(package["points"], "cumulative_planned_pct")
            for package in packages
        ) / total_weight
        actual = sum(
            float(package["weight"] or 0) * value_at(package["points"], "cumulative_actual_pct")
            for package in packages
        ) / total_weight
        result[int(scheme_id)] = {
            "month": cutoff,
            "planned": round(planned, 2),
            "actual": round(actual, 2),
            "variance": round(actual - planned, 2),
        }
    return result


def _build_live_preview(report_id: str, fy: str, db: Session) -> dict:
    fy = _normalize_fy(fy)
    capex = _capex_project_rows(db, fy)
    generated = datetime.now().strftime("%Y-%m-%d %H:%M")

    if report_id in {"mos-capex", "capex-pmc", "capex-recon"}:
        capex_report = _capex_report_rows(db, fy)
        title = {
            "mos-capex": "MoS CAPEX Format",
            "capex-pmc": "CAPEX PMC Report",
            "capex-recon": "CAPEX Reconciliation",
        }[report_id]
        columns = ["Sl", "Scheme", "Gross Cost", "Cum till last FY", f"BE {fy}", f"RE {fy}", "Actual YTD", "Achievement %"]
        rows = []
        for index, item in enumerate(capex_report, 1):
            be = float(item["be_fy"] or 0)
            actual = float(item["actual_fy"] or 0)
            rows.append([
                index, item["scheme_name"], round(float(item["gross_cost"] or 0), 2),
                round(float(item["cum_last_fy"] or 0), 2), round(be, 2),
                round(float(item["re_fy"] or 0), 2), round(actual, 2),
                round(actual / be * 100, 2) if be else 0,
            ])
        totals = [
            "", "TOTAL",
            round(sum(float(item["gross_cost"] or 0) for item in capex_report), 2),
            round(sum(float(item["cum_last_fy"] or 0) for item in capex_report), 2),
            round(sum(float(item["be_fy"] or 0) for item in capex_report), 2),
            round(sum(float(item["re_fy"] or 0) for item in capex_report), 2),
            round(sum(float(item["actual_fy"] or 0) for item in capex_report), 2),
            0,
        ]
        totals[-1] = round(totals[6] / totals[4] * 100, 2) if totals[4] else 0
        rows.append(totals)
        return {"title": title, "fy": fy, "generated": generated, "columns": columns, "rows": rows,
                "footnote": "Live CAPEX plan; parent projects are excluded from leaf-project totals."}

    if report_id in {"phys-fin", "pmc-phys"}:
        physical = _physical_progress_by_scheme(db)
        columns = ["Sl", "Scheme", "Gross Cost", "Cumulative Exp", "Financial %", "Physical Plan %", "Physical Actual %", "Variance", "Status"]
        rows = []
        for index, item in enumerate(capex, 1):
            gross = float(item["gross_cost"] or 0)
            cumulative = float(item["cum_last_fy"] or 0) + float(item["actual_fy"] or 0)
            progress = physical.get(int(item["scheme_id"]), {})
            variance = float(progress.get("variance", 0))
            status = "On Track" if variance >= 0 else "At Risk" if variance >= -5 else "Behind"
            rows.append([
                index, item["scheme_name"], round(gross, 2), round(cumulative, 2),
                round(cumulative / gross * 100, 2) if gross else 0,
                progress.get("planned", 0), progress.get("actual", 0), round(variance, 2), status,
            ])
        return {"title": "Physical Progress — PMC" if report_id == "pmc-phys" else "Physical & Financial Progress",
                "fy": fy, "generated": generated, "columns": columns, "rows": rows,
                "footnote": "Physical progress uses current locked plans and activity weightages; financial progress uses CAPEX actuals."}

    if report_id == "s-curve":
        from app.api.v1.s_curve import get_s_curve_data
        scheme_rows = db.execute(text("""
            SELECT DISTINCT s.scheme_id, COALESCE(s.estimated_cost_cr, 1) AS weight
            FROM scheme_master s
            JOIN packages p ON p.scheme_id = s.scheme_id AND NOT p.is_deleted
            JOIN progress_plans pp ON pp.package_id = p.package_id
            WHERE pp.is_current AND pp.is_locked AND NOT pp.is_deleted AND NOT s.is_deleted
        """)).mappings().all()
        curves = []
        for scheme in scheme_rows:
            curve = get_s_curve_data(int(scheme["scheme_id"]), db)
            if curve.get("planned"):
                curves.append((float(scheme["weight"] or 1), curve))
        month_order = ["Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec", "Jan", "Feb", "Mar"]
        months = []
        for _weight, curve in curves:
            for point in curve["planned"]:
                if point["month"] not in months:
                    months.append(point["month"])
        months.sort(key=lambda month: (month_order.index(month[:3]), month))
        total_weight = sum(weight for weight, _curve in curves) or 1
        rows = []
        previous_plan = previous_actual = 0.0
        for month in months:
            planned = actual = 0.0
            actual_weight = 0.0
            for weight, curve in curves:
                plan_map = {point["month"]: float(point["value"]) for point in curve["planned"]}
                actual_map = {point["month"]: float(point["value"]) for point in curve["actual"]}
                planned += weight * plan_map.get(month, 0)
                if month in actual_map:
                    actual += weight * actual_map[month]
                    actual_weight += weight
            cumulative_plan = planned / total_weight
            cumulative_actual = actual / actual_weight if actual_weight else None
            rows.append([
                month, round(cumulative_plan - previous_plan, 2),
                round(cumulative_actual - previous_actual, 2) if cumulative_actual is not None else "",
                round(cumulative_plan, 2), round(cumulative_actual, 2) if cumulative_actual is not None else "",
                round(cumulative_actual - cumulative_plan, 2) if cumulative_actual is not None else "",
            ])
            previous_plan = cumulative_plan
            if cumulative_actual is not None:
                previous_actual = cumulative_actual
        return {"title": "Weighted S-Curve — Portfolio", "fy": fy, "generated": generated,
                "columns": ["Month", "Monthly Plan %", "Monthly Actual %", "Cum Plan %", "Cum Actual %", "Variance"],
                "rows": rows, "footnote": "Live current locked plans; activity and package weightages applied."}

    if report_id == "dpr":
        physical = _physical_progress_by_scheme(db)
        names = {
            row["scheme_id"]: row["scheme_name"]
            for row in db.execute(text("SELECT scheme_id, scheme_name FROM scheme_master WHERE NOT is_deleted")).mappings()
        }
        rows = []
        for index, (scheme_id, progress) in enumerate(sorted(physical.items()), 1):
            rows.append([
                index, names.get(scheme_id, f"Scheme {scheme_id}"), progress["month"],
                progress["planned"], progress["actual"], progress["variance"],
                "On Track" if progress["variance"] >= 0 else "At Risk" if progress["variance"] >= -5 else "Behind",
            ])
        return {"title": "DPR Summary", "fy": fy, "generated": generated,
                "columns": ["Sl", "Scheme", "Cut-off", "Planned %", "Actual %", "Variance", "Status"],
                "rows": rows, "footnote": "Live current-plan DPR rollup at the latest actual month."}

    raise HTTPException(status_code=404, detail=f"Live preview not configured for {report_id}")


@router.get("/preview")
def preview_live_report(
    id: str = Query(...),
    fy: str = Query("2026-27"),
    db: Session = Depends(get_db),
):
    return _build_live_preview(id, fy, db)


@router.get("/export")
def export_live_report(
    id: str = Query(...),
    fy: str = Query("2026-27"),
    fmt: str = Query("csv"),
    db: Session = Depends(get_db),
):
    preview = _build_live_preview(id, fy, db)
    safe_name = f"{id}-{_normalize_fy(fy)}"
    fmt = fmt.lower()
    if fmt == "csv":
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(preview["columns"])
        writer.writerows(preview["rows"])
        return StreamingResponse(
            iter([output.getvalue()]), media_type="text/csv",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.csv"'},
        )
    if fmt == "xlsx":
        from openpyxl import Workbook
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Report"
        sheet.append(preview["columns"])
        for row in preview["rows"]:
            sheet.append(row)
        output = io.BytesIO()
        workbook.save(output)
        output.seek(0)
        return StreamingResponse(output, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                 headers={"Content-Disposition": f'attachment; filename="{safe_name}.xlsx"'})
    if fmt == "docx":
        from docx import Document
        document = Document()
        document.add_heading(preview["title"], level=1)
        table = document.add_table(rows=1, cols=len(preview["columns"]))
        for index, column in enumerate(preview["columns"]):
            table.rows[0].cells[index].text = str(column)
        for row in preview["rows"]:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = str(value)
        output = io.BytesIO()
        document.save(output)
        output.seek(0)
        return StreamingResponse(output, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                 headers={"Content-Disposition": f'attachment; filename="{safe_name}.docx"'})
    if fmt == "pdf":
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        output = io.BytesIO()
        document = SimpleDocTemplate(output, pagesize=landscape(A4))
        styles = getSampleStyleSheet()
        table_data = [preview["columns"], *preview["rows"]]
        table = Table(table_data, repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbeafe")),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        document.build([Paragraph(preview["title"], styles["Title"]), Spacer(1, 8), table])
        output.seek(0)
        return StreamingResponse(output, media_type="application/pdf",
                                 headers={"Content-Disposition": f'attachment; filename="{safe_name}.pdf"'})
    raise HTTPException(status_code=400, detail="fmt must be csv, xlsx, docx, or pdf")


# ============================================================================
# 1) POST /upload  →  upload a .docx, convert to HTML, store in DB
# ============================================================================
@router.post("/upload")
async def upload_report_doc(
    title: str = "Untitled Report",
    scheme_id: int = 74,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    if not file.filename or not file.filename.lower().endswith((".docx", ".doc")):
        raise HTTPException(status_code=400, detail="Only .docx files accepted")

    raw = await file.read()

    # Convert .docx → HTML via mammoth
    try:
        import mammoth
        result = mammoth.convert_to_html(io.BytesIO(raw))
        html = result.value
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Could not parse .docx: {e}")

    # Store the original .docx bytes as base64 in extra_fields for re-export
    import base64
    docx_b64 = base64.b64encode(raw).decode("ascii")

    # Create a slug from the filename
    slug = file.filename.rsplit(".", 1)[0].lower().replace(" ", "-").replace("_", "-")

    # Upsert: if a doc with this slug already exists, update it
    existing = db.execute(text("""
        SELECT note_id FROM record_notes
        WHERE note_type='report_doc' AND extra_fields->>'doc_key' = :k AND is_deleted=FALSE
        LIMIT 1
    """), {"k": slug}).first()

    ef = json.dumps({
        "doc_key": slug,
        "original_filename": file.filename,
        "original_docx_b64": docx_b64,
    })

    if existing:
        db.execute(text("""
            UPDATE record_notes
            SET title = :t, body = :b, extra_fields = CAST(:ef AS jsonb),
                updated_at = CURRENT_TIMESTAMP
            WHERE note_id = :id
        """), {"t": title, "b": html, "ef": ef, "id": existing.note_id})
        note_id = existing.note_id
    else:
        note_id = db.execute(text("""
            INSERT INTO record_notes (
                scheme_id, note_type, title, body, extra_fields,
                is_deleted, created_by, created_at, updated_at
            ) VALUES (
                :sid, 'report_doc', :t, :b, CAST(:ef AS jsonb),
                FALSE, 1, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP
            ) RETURNING note_id
        """), {"sid": scheme_id, "t": title, "b": html, "ef": ef}).scalar()

    db.commit()
    return {
        "ok": True,
        "note_id": note_id,
        "slug": slug,
        "title": title,
        "html_length": len(html),
    }


# ============================================================================
# 2) GET /list  →  all report docs
# ============================================================================
@router.get("/list")
def list_report_docs(db: Session = Depends(get_db)):
    rows = db.execute(text("""
        SELECT note_id, title,
               extra_fields->>'doc_key' AS slug,
               extra_fields->>'original_filename' AS filename,
               created_at, updated_at
        FROM record_notes
        WHERE note_type = 'report_doc' AND is_deleted = FALSE
        ORDER BY updated_at DESC NULLS LAST
    """)).fetchall()
    return [{
        "note_id": r.note_id,
        "title": r.title,
        "slug": r.slug,
        "filename": r.filename,
        "created_at": r.created_at.isoformat() if r.created_at else None,
        "updated_at": r.updated_at.isoformat() if r.updated_at else None,
    } for r in rows]


# ============================================================================
# 3) GET /{slug}  →  get one doc's HTML
# ============================================================================
@router.get("/{slug}")
def get_report_doc(slug: str, db: Session = Depends(get_db)):
    row = db.execute(text("""
        SELECT note_id, title, body, extra_fields, updated_at
        FROM record_notes
        WHERE note_type='report_doc' AND extra_fields->>'doc_key' = :k AND is_deleted=FALSE
        ORDER BY updated_at DESC NULLS LAST LIMIT 1
    """), {"k": slug}).first()
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")
    ef = row.extra_fields or {}
    return {
        "note_id": row.note_id,
        "title": row.title,
        "slug": slug,
        "html": row.body,
        "filename": ef.get("original_filename"),
        "updated_at": row.updated_at.isoformat() if row.updated_at else None,
    }


# ============================================================================
# 4) PUT /{slug}  →  save edited HTML
# ============================================================================
@router.put("/{slug}")
def save_report_doc(slug: str, payload: dict = Body(...), db: Session = Depends(get_db)):
    html = payload.get("html", "")
    title = payload.get("title")
    if not html.strip():
        raise HTTPException(status_code=400, detail="Empty document")

    existing = db.execute(text("""
        SELECT note_id FROM record_notes
        WHERE note_type='report_doc' AND extra_fields->>'doc_key' = :k AND is_deleted=FALSE
        LIMIT 1
    """), {"k": slug}).first()
    if not existing:
        raise HTTPException(status_code=404, detail="Document not found")

    sets = ["body = :b", "updated_at = CURRENT_TIMESTAMP"]
    params = {"b": html, "id": existing.note_id}
    if title:
        sets.append("title = :t")
        params["t"] = title

    db.execute(text(f"UPDATE record_notes SET {', '.join(sets)} WHERE note_id = :id"), params)
    db.commit()
    return {"ok": True, "updated_at": datetime.utcnow().isoformat()}


# ============================================================================
# 5) GET /{slug}/export?format=docx|pdf  →  download
# ============================================================================
@router.get("/{slug}/export")
def export_report_doc(slug: str, format: str = "docx", db: Session = Depends(get_db)):
    row = db.execute(text("""
        SELECT title, body, extra_fields FROM record_notes
        WHERE note_type='report_doc' AND extra_fields->>'doc_key' = :k AND is_deleted=FALSE
        LIMIT 1
    """), {"k": slug}).first()
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")

    title = row.title or slug
    html = row.body or ""

    if format == "docx":
        return _export_docx(html, title)
    elif format == "pdf":
        return _export_pdf(html, title)
    else:
        raise HTTPException(status_code=400, detail="format must be 'docx' or 'pdf'")


def _export_docx(html: str, title: str):
    try:
        from docx import Document
        from htmldocx import HtmlToDocx
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"DOCX export dependency missing. Install: python-docx htmldocx ({e})",
        )

    doc = Document()
    # Set narrow margins for A4
    for section in doc.sections:
        section.top_margin = 914400     # 1 inch
        section.bottom_margin = 914400
        section.left_margin = 914400
        section.right_margin = 914400

    parser = HtmlToDocx()
    parser.add_html_to_document(html, doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = title.replace(" ", "_")[:50] + ".docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}"'},
    )


def _export_pdf(html: str, title: str):
    """Simple PDF export via reportlab. For complex HTML, weasyprint is better
    but requires system-level install. This gives a clean basic PDF."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import mm
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"PDF export dependency missing. Install: reportlab ({e})",
        )
    import re

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=18*mm, rightMargin=18*mm,
                            topMargin=20*mm, bottomMargin=20*mm)
    styles = getSampleStyleSheet()
    story = []

    # Strip HTML to paragraphs (simple approach — good for text-heavy docs)
    # For full HTML→PDF fidelity, install weasyprint
    text_content = re.sub(r'<br\s*/?>', '\n', html)
    text_content = re.sub(r'<[^>]+>', '', text_content)
    for line in text_content.split('\n'):
        line = line.strip()
        if line:
            story.append(Paragraph(line, styles['Normal']))
        else:
            story.append(Spacer(1, 4*mm))

    doc.build(story)
    buf.seek(0)
    safe_name = title.replace(" ", "_")[:50] + ".pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}"'},
    )


# ============================================================================
# 6) DELETE /{slug}  →  soft-delete
# ============================================================================
@router.delete("/{slug}")
def delete_report_doc(slug: str, db: Session = Depends(get_db)):
    db.execute(text("""
        UPDATE record_notes SET is_deleted = TRUE
        WHERE note_type='report_doc' AND extra_fields->>'doc_key' = :k
    """), {"k": slug})
    db.commit()
    return {"ok": True}
