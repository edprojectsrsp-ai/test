"""Report Studio API — self-service KPI / metric builder.

  GET    /report-studio/datasets            curated dataset registry (fields + measures)
  POST   /report-studio/query               run an ad-hoc structured query
  GET    /report-studio/metrics             list saved metrics
  POST   /report-studio/metrics             save a metric (query spec + viz)
  GET    /report-studio/metrics/{id}        fetch one
  PUT    /report-studio/metrics/{id}        update
  DELETE /report-studio/metrics/{id}        delete
  POST   /report-studio/metrics/{id}/run    execute a saved metric

All queries are compiled to safe parameterized SQL against the dataset registry
(app/services/report_studio.py). No raw user SQL is accepted or executed.
"""
from __future__ import annotations

import json
from typing import Any, Optional

from fastapi import APIRouter, Depends, HTTPException
from pydantic import BaseModel
from sqlalchemy import text
from sqlalchemy.orm import Session

from app.core.database import get_db
from app.security.auth import require_user
from app.services import report_studio as RS

router = APIRouter(
    prefix="/report-studio",
    tags=["Report Studio"],
    dependencies=[Depends(require_user)],
)


def _run_query(db: Session, q: RS.QueryIn) -> dict[str, Any]:
    try:
        sql, params, columns = RS.compile_query(q)
    except RS.CompileError as e:
        raise HTTPException(status_code=400, detail=str(e))
    try:
        rows = db.execute(text(sql), params).mappings().all()
    except Exception as e:  # surface SQL errors as 400 (bad query), not 500
        raise HTTPException(status_code=400, detail=f"Query failed: {str(e)[:300]}")
    # jsonify (dates, Decimals) — mappings() already gives dict-likes; coerce values
    out_rows = []
    for r in rows:
        row = {}
        for k, v in dict(r).items():
            if hasattr(v, "isoformat"):
                row[k] = v.isoformat()
            elif isinstance(v, (int, float, str, bool)) or v is None:
                row[k] = v
            else:
                row[k] = float(v)  # Decimal
        out_rows.append(row)
    columns, out_rows = RS.apply_postprocess(q, columns, out_rows)
    return {"columns": columns, "rows": out_rows, "sql": sql, "row_count": len(out_rows)}


@router.get("/datasets")
def datasets():
    return {"datasets": RS.registry_public()}


@router.post("/query")
def run_query(q: RS.QueryIn, db: Session = Depends(get_db)):
    return _run_query(db, q)


@router.get("/field-values")
def field_values(dataset: str, field: str, search: Optional[str] = None,
                 limit: int = 200, db: Session = Depends(get_db)):
    """Distinct values of a dimension — powers the filter member-picker."""
    try:
        sql, params = RS.compile_field_values(dataset, field, search, limit)
    except RS.CompileError as e:
        raise HTTPException(status_code=400, detail=str(e))
    try:
        rows = db.execute(text(sql), params).all()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Lookup failed: {str(e)[:200]}")
    vals = []
    for (v,) in rows:
        if v is None:
            continue
        vals.append(v.isoformat() if hasattr(v, "isoformat") else v)
    return {"dataset": dataset, "field": field, "values": vals}


# ---------------------------------------------------------------- saved metrics

class MetricIn(BaseModel):
    name: str
    description: Optional[str] = None
    dataset: str
    spec: RS.QueryIn
    viz: str = "kpi"
    folder: Optional[str] = None
    is_pinned: bool = False


def _metric_row(db: Session, metric_id: int) -> dict:
    row = db.execute(
        text("SELECT * FROM rs_metrics WHERE metric_id = :m"), {"m": metric_id}
    ).mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Metric not found")
    return dict(row)


@router.get("/metrics")
def list_metrics(db: Session = Depends(get_db)):
    rows = db.execute(text(
        "SELECT metric_id, name, description, dataset, viz, folder, is_pinned, updated_at "
        "FROM rs_metrics ORDER BY is_pinned DESC, updated_at DESC"
    )).mappings().all()
    return {"metrics": [dict(r) for r in rows]}


@router.post("/metrics")
def create_metric(payload: MetricIn, db: Session = Depends(get_db)):
    if payload.dataset not in RS.DATASETS:
        raise HTTPException(status_code=400, detail=f"Unknown dataset '{payload.dataset}'")
    # validate the spec compiles before saving
    try:
        RS.compile_query(payload.spec)
    except RS.CompileError as e:
        raise HTTPException(status_code=400, detail=f"Invalid metric spec: {e}")
    mid = db.execute(text(
        "INSERT INTO rs_metrics (name, description, dataset, spec, viz, folder, is_pinned) "
        "VALUES (:n, :d, :ds, CAST(:spec AS jsonb), :viz, :folder, :pin) RETURNING metric_id"
    ), {
        "n": payload.name, "d": payload.description, "ds": payload.dataset,
        "spec": payload.spec.model_dump_json(), "viz": payload.viz,
        "folder": payload.folder, "pin": payload.is_pinned,
    }).scalar()
    db.commit()
    return {"metric_id": mid}


@router.get("/metrics/{metric_id}")
def get_metric(metric_id: int, db: Session = Depends(get_db)):
    row = _metric_row(db, metric_id)
    if isinstance(row.get("spec"), str):
        row["spec"] = json.loads(row["spec"])
    return row


@router.put("/metrics/{metric_id}")
def update_metric(metric_id: int, payload: MetricIn, db: Session = Depends(get_db)):
    _metric_row(db, metric_id)
    try:
        RS.compile_query(payload.spec)
    except RS.CompileError as e:
        raise HTTPException(status_code=400, detail=f"Invalid metric spec: {e}")
    db.execute(text(
        "UPDATE rs_metrics SET name=:n, description=:d, dataset=:ds, spec=CAST(:spec AS jsonb), "
        "viz=:viz, folder=:folder, is_pinned=:pin, updated_at=now() WHERE metric_id=:m"
    ), {
        "n": payload.name, "d": payload.description, "ds": payload.dataset,
        "spec": payload.spec.model_dump_json(), "viz": payload.viz,
        "folder": payload.folder, "pin": payload.is_pinned, "m": metric_id,
    })
    db.commit()
    return {"ok": True}


@router.delete("/metrics/{metric_id}")
def delete_metric(metric_id: int, db: Session = Depends(get_db)):
    _metric_row(db, metric_id)
    db.execute(text("DELETE FROM rs_metrics WHERE metric_id = :m"), {"m": metric_id})
    db.commit()
    return {"ok": True}


@router.post("/metrics/{metric_id}/run")
def run_metric(metric_id: int, db: Session = Depends(get_db)):
    row = _metric_row(db, metric_id)
    spec = row["spec"]
    if isinstance(spec, str):
        spec = json.loads(spec)
    q = RS.QueryIn(**spec)
    result = _run_query(db, q)
    result["metric"] = {"metric_id": metric_id, "name": row["name"], "viz": row["viz"]}
    return result


# ---------------------------------------------------------------- custom reports
# A custom report = an ordered list of sections, each a full Report Studio
# query spec (dimensions/measures/formulas/filters/pivot/totals) + a title.
# Runs live against the semantic layer; exports to XLSX and DOCX.

class SectionIn(BaseModel):
    title: str
    note: Optional[str] = None
    spec: RS.QueryIn


class ReportIn(BaseModel):
    name: str
    description: Optional[str] = None
    category: Optional[str] = None
    sections: list[SectionIn]


def _ensure_reports_table(db: Session):
    db.execute(text(
        "CREATE TABLE IF NOT EXISTS rs_reports ("
        " report_id SERIAL PRIMARY KEY,"
        " name VARCHAR(200) NOT NULL,"
        " description TEXT,"
        " category VARCHAR(80),"
        " sections JSONB NOT NULL DEFAULT '[]'::jsonb,"
        " created_at TIMESTAMP NOT NULL DEFAULT now(),"
        " updated_at TIMESTAMP NOT NULL DEFAULT now())"
    ))
    db.commit()


def _validate_sections(payload: ReportIn):
    if not payload.sections:
        raise HTTPException(status_code=400, detail="A report needs at least one section")
    for s in payload.sections:
        try:
            RS.compile_query(s.spec)
        except RS.CompileError as e:
            raise HTTPException(status_code=400, detail=f"Section '{s.title}': {e}")


def _report_row(db: Session, report_id: int) -> dict:
    row = db.execute(text("SELECT * FROM rs_reports WHERE report_id = :r"),
                     {"r": report_id}).mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Report not found")
    out = dict(row)
    if isinstance(out.get("sections"), str):
        out["sections"] = json.loads(out["sections"])
    return out


@router.get("/reports")
def list_reports(category: Optional[str] = None, db: Session = Depends(get_db)):
    _ensure_reports_table(db)
    where = "WHERE category = :cat" if category else ""
    rows = db.execute(text(
        "SELECT report_id, name, description, category, "
        "       jsonb_array_length(sections) AS section_count, updated_at "
        f"FROM rs_reports {where} ORDER BY updated_at DESC"
    ), {"cat": category} if category else {}).mappings().all()
    return {"reports": [dict(r) for r in rows]}


@router.post("/reports")
def create_report(payload: ReportIn, db: Session = Depends(get_db)):
    _ensure_reports_table(db)
    _validate_sections(payload)
    rid = db.execute(text(
        "INSERT INTO rs_reports (name, description, category, sections) "
        "VALUES (:n, :d, :c, CAST(:s AS jsonb)) RETURNING report_id"
    ), {"n": payload.name, "d": payload.description, "c": payload.category,
        "s": json.dumps([s.model_dump(mode="json") for s in payload.sections])}).scalar()
    db.commit()
    return {"report_id": rid}


@router.get("/reports/{report_id}")
def get_report(report_id: int, db: Session = Depends(get_db)):
    _ensure_reports_table(db)
    return _report_row(db, report_id)


@router.put("/reports/{report_id}")
def update_report(report_id: int, payload: ReportIn, db: Session = Depends(get_db)):
    _ensure_reports_table(db)
    _report_row(db, report_id)
    _validate_sections(payload)
    db.execute(text(
        "UPDATE rs_reports SET name=:n, description=:d, category=:c, "
        "sections=CAST(:s AS jsonb), updated_at=now() WHERE report_id=:r"
    ), {"n": payload.name, "d": payload.description, "c": payload.category,
        "s": json.dumps([s.model_dump(mode="json") for s in payload.sections]),
        "r": report_id})
    db.commit()
    return {"ok": True}


@router.delete("/reports/{report_id}")
def delete_report(report_id: int, db: Session = Depends(get_db)):
    _ensure_reports_table(db)
    _report_row(db, report_id)
    db.execute(text("DELETE FROM rs_reports WHERE report_id = :r"), {"r": report_id})
    db.commit()
    return {"ok": True}


def _run_report(db: Session, report: dict) -> dict[str, Any]:
    sections_out = []
    for s in report["sections"]:
        q = RS.QueryIn(**s["spec"])
        res = _run_query(db, q)
        sections_out.append({
            "title": s.get("title") or "Section",
            "note": s.get("note"),
            "columns": res["columns"],
            "rows": res["rows"],
        })
    return {
        "report_id": report["report_id"], "name": report["name"],
        "description": report.get("description"), "category": report.get("category"),
        "sections": sections_out,
    }


@router.post("/reports/{report_id}/run")
def run_report(report_id: int, db: Session = Depends(get_db)):
    _ensure_reports_table(db)
    return _run_report(db, _report_row(db, report_id))


# ---------------------------------------------------------------- export

def _fmt_cell(v: Any) -> Any:
    if isinstance(v, float):
        return round(v, 2)
    return v


@router.get("/reports/{report_id}/export")
def export_report(report_id: int, fmt: str = "xlsx", db: Session = Depends(get_db)):
    import io

    from fastapi.responses import StreamingResponse

    _ensure_reports_table(db)
    data = _run_report(db, _report_row(db, report_id))
    stem = "".join(c if c.isalnum() or c in " -_" else "_" for c in data["name"]).strip()[:80] or "report"

    if fmt == "xlsx":
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

        wb = Workbook()
        wb.remove(wb.active)
        thin = Side(style="thin", color="999999")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)
        head_fill = PatternFill("solid", fgColor="0B3D91")
        head_font = Font(bold=True, color="FFFFFF", size=10)
        title_font = Font(bold=True, size=13, color="0B3D91")
        used_names: set[str] = set()
        for i, sec in enumerate(data["sections"], 1):
            base = "".join(c for c in sec["title"] if c not in "[]:*?/\\")[:28] or f"Section {i}"
            name = base
            n = 1
            while name in used_names:
                n += 1
                name = f"{base[:25]} {n}"
            used_names.add(name)
            ws = wb.create_sheet(title=name)
            cols = [c for c in sec["columns"] if not c["key"].startswith("__")]
            ws.cell(row=1, column=1, value=sec["title"]).font = title_font
            if sec.get("note"):
                ws.cell(row=2, column=1, value=sec["note"]).font = Font(italic=True, size=9, color="666666")
            hr = 3
            for ci, c in enumerate(cols, 1):
                cell = ws.cell(row=hr, column=ci, value=c["label"])
                cell.font = head_font
                cell.fill = head_fill
                cell.border = border
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                ws.column_dimensions[cell.column_letter].width = max(11, min(38, len(c["label"]) + 4))
            for ri, r in enumerate(sec["rows"], hr + 1):
                is_total = bool(r.get("__total__"))
                for ci, c in enumerate(cols, 1):
                    cell = ws.cell(row=ri, column=ci, value=_fmt_cell(r.get(c["key"])))
                    cell.border = border
                    if c["type"] in ("int", "number", "money"):
                        cell.number_format = "#,##0.00" if c["type"] != "int" else "#,##0"
                    if is_total:
                        cell.font = Font(bold=True)
                        cell.fill = PatternFill("solid", fgColor="E8EEF9")
            ws.freeze_panes = ws.cell(row=hr + 1, column=2)
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f'attachment; filename="{stem}.xlsx"'})

    if fmt == "docx":
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        doc = Document()
        sec0 = doc.sections[0]
        sec0.orientation = WD_ORIENT.LANDSCAPE
        sec0.page_width, sec0.page_height = sec0.page_height, sec0.page_width
        h = doc.add_heading(data["name"], level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if data.get("description"):
            p = doc.add_paragraph(data["description"])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for sec in data["sections"]:
            doc.add_heading(sec["title"], level=2)
            if sec.get("note"):
                doc.add_paragraph(sec["note"]).runs[0].font.size = Pt(8)
            cols = [c for c in sec["columns"] if not c["key"].startswith("__")]
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for ci, c in enumerate(cols):
                hdr[ci].text = c["label"]
                for run in hdr[ci].paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(8)
            for r in sec["rows"]:
                cells = table.add_row().cells
                for ci, c in enumerate(cols):
                    v = _fmt_cell(r.get(c["key"]))
                    cells[ci].text = "" if v is None else (f"{v:,.2f}" if isinstance(v, float) else str(v))
                    for run in cells[ci].paragraphs[0].runs:
                        run.font.size = Pt(8)
                        if r.get("__total__"):
                            run.font.bold = True
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{stem}.docx"'})

    raise HTTPException(status_code=400, detail="fmt must be xlsx or docx")


# ---------------------------------------------------------------- seed: CAPEX pack

CAPEX_PACK = "capex-pack"


MOS_CAPEX_SOURCE_ROWS = [
    ("r6", 6, "", "Ontime", 9, 5368.3087065842, 1008.39, 1357, 611.24, 1619.63),
    ("r7", 7, "", "Delay<1", 1, 0, 0, 0, 0, 0),
    ("r8", 8, "", "Delay>1", 0, 0, 0, 0, 0, 0),
    ("r10", 10, "", "Ontime", 5, 70.07738763, 0.1682035, 29, 0.7443474, 0.9125509),
    ("r11", 11, "", "Delay<1", 9, 116.252210912, 37.735304047, 63, 10.9388835, 48.674187547),
    ("r12", 12, "", "Delay>1", 4, 73.911584264, 32.47813728, 34, 7.0826475, 39.56078478),
    ("r15", 15, "", "Ontime", 0, 0, 0, 0, 0, 0),
    ("r16", 16, "", "Delay<1", 0, 0, 0, 0, 0, 0),
    ("r17", 17, "", "Delay>1", 0, 0, 0, 0, 0, 0),
    ("r19", 19, "", "Ontime", 23, 183.4184411419, 0, 33, 0.0093594, 0.0093594),
    ("r20", 20, "", "Delay<1", 1, 0.3855296, 0, 0.3, 0, 0),
    ("r21", 21, "", "Delay>1", 0, 0, 0, 0, 0, 0),
    ("r23", 23, "", "Ontime", 37, 5621.8045353561, 1008.5582035, 1419, 611.9937068, 1620.5519103),
    ("r24", 24, "", "Delay<1", 11, 116.637740512, 37.735304047, 63.3, 10.9388835, 48.674187547),
    ("r25", 25, "", "Delay>1", 4, 73.911584264, 32.47813728, 34, 7.0826475, 39.56078478),
    ("r26", 26, "2", "Milestone payments in completed projects", None, None, None, 288.5081145, 172.3147622, 172.3147622),
    ("r27", 27, "", "MEP", 1, 14240, 14183.99, 0, 0, 14183.99),
    ("r28", 28, "", "Others (COB#4)", None, None, None, None, 0, 0),
    ("r29", 29, "", "30 MLD STP for township", 1, 96.62, 87.02, 12.2, 0.87, 87.89),
    ("r30", 30, "", "TS#1", 1, 158.17, 165.65, 0, 0.21, 165.86),
    ("r31", 31, "", "COB #2", 1, 433.58, 461.48, 5, 1.17, 462.65),
    ("r32", 32, "", "Caster-IV", 1, 1105.23, 897.96, 195, 77.7, 975.66),
    ("r33", 33, "", "NHSM", 1, 4033.88, 4177.77, 0, 0, 4177.77),
    ("r34", 34, "", "Power enabling package for oxygen", 1, 83.49, 60.28, 16, 5.29, 65.57),
    ("r35", 35, "", "Plant AMR(<30 Cr.)-Completed/EDC", None, None, None, 34.59, 80.574109409, 80.574109409),
    ("r36", 36, "", "Plant AMR(<30 Cr.)-old schemes", 9, 49.970598566, 14.44075282, 24.38, 5.662538291, 20.103291111),
    ("r37", 37, "", "Plant AMR(<30 Cr.)-new schemes", 2, 1.4124068, 0, 1.3381145, 0.8381145, 0.8381145),
    ("r39", 39, "", "Corporate AMR-under tendering", 1, 367.15, 0.74, 0, 0, 0.74),
    ("r40", 40, "", "Corporate AMR- under final approval and contract award", 0, 0, 0, 0, 0, 0),
    ("r41", 41, "", "Plant AMR(<30 Cr.)", 14, 72.976383703, 0, 10.4, 0, 0),
    ("r43", 43, "", "Dust Extraction for BF#1", 1, 52.56, 0, 0, 0, 0),
    ("r44", 44, "", "Provision of Mixed Gas Firing in oil fired kilns at CP II", 1, 64.66, 0, 0, 0, 0),
    ("r45", 45, "", "Dust Extraction for SP-III", 1, 70.71, 0, 0, 0, 0),
    ("r46", 46, "", "Upgradation of Electrics for 2 nos. Hot Metal Handling Cranes at SMS-II", 1, 61.94, 0, 0, 0, 0),
    ("r47", 47, "", "Others-CET Consultancy debit, etc", None, None, None, 0, 0, 0),
    ("r49", 49, "4", "Spares & Capital Repairs", None, None, None, 334.8, 2.2, 2.2),
    ("r50", 50, "5", "Other schemes/ JVs", None, None, None, None, None, 0),
]


DRILLDOWN_MAP = {
    "1a": ["r6", "r7", "r8", "r10", "r11", "r12"],
    "1b": ["r15", "r16", "r17", "r19", "r20", "r21"],
    "1": ["r23", "r24", "r25"],
    "2": ["r27", "r28", "r29", "r30", "r31", "r32", "r33", "r34", "r35", "r36", "r37"],
    "3a": ["r39", "r40", "r41"],
    "3b": ["r43", "r44", "r45", "r46", "r47"],
    "3": ["r39", "r40", "r41", "r43", "r44", "r45", "r46", "r47"],
    "4": ["r49"],
    "5": ["r50"],
    "total": ["r23", "r24", "r25", "r27", "r28", "r29", "r30", "r31", "r32", "r33", "r34", "r35", "r36", "r37", "r39", "r40", "r41", "r43", "r44", "r45", "r46", "r47", "r49"],
}


def _ensure_mos_capex_source_rows(db: Session):
    db.execute(text("""
        CREATE TABLE IF NOT EXISTS mos_capex_source_rows (
            report_key TEXT NOT NULL,
            row_key TEXT NOT NULL,
            source_row INTEGER NOT NULL,
            sn TEXT,
            category TEXT NOT NULL,
            project_count NUMERIC,
            total_cost NUMERIC,
            exp_upto_last_fy NUMERIC,
            capex_re NUMERIC,
            exp_current_fy NUMERIC,
            total_exp NUMERIC,
            source_file TEXT,
            source_sheet TEXT,
            PRIMARY KEY (report_key, row_key)
        )
    """))
    db.execute(text("DELETE FROM mos_capex_source_rows WHERE report_key = 'fy25_26_mos_capex'"))
    db.execute(text("""
        INSERT INTO mos_capex_source_rows
            (report_key, row_key, source_row, sn, category, project_count, total_cost,
             exp_upto_last_fy, capex_re, exp_current_fy, total_exp, source_file, source_sheet)
        VALUES
            ('fy25_26_mos_capex', :row_key, :source_row, :sn, :category, :project_count,
             :total_cost, :exp_upto_last_fy, :capex_re, :exp_current_fy, :total_exp,
             '1.1,2,4 RSP MoS presentation-FY 26-27 Backup03.07.2026-PMC- Ongoing -R.xlsx',
             'MOS Capex Format')
    """), [
        {
            "row_key": row_key,
            "source_row": source_row,
            "sn": sn,
            "category": category,
            "project_count": project_count,
            "total_cost": total_cost,
            "exp_upto_last_fy": exp_upto_last_fy,
            "capex_re": capex_re,
            "exp_current_fy": exp_current_fy,
            "total_exp": total_exp,
        }
        for row_key, source_row, sn, category, project_count, total_cost,
        exp_upto_last_fy, capex_re, exp_current_fy, total_exp in MOS_CAPEX_SOURCE_ROWS
    ])
    db.commit()


MOS_RECON_ROW_MAP = {
    "1a": ("Being implemented from last FY", "1"),
    "1b": ("Implementation started during FY24-25", "2"),
    "1": ("Total Ongoing projects (1a+1b)", "3"),
    "2": ("Milestone payments in completed projects", "milestone"),
    "3a": ("New Projects under tendering/ final approval and contract award", "3a"),
    "3b": ("New Projects under Stage-I approval", "3b"),
    "3": ("Total New projects under consideration (3a+3b)", "total_new"),
    "4": ("Spares & Capital Repairs", "spares"),
    "5": ("Other schemes/ JVs", "other"),
    "total": ("Total", "total"),
}


def _source_summary_rows(db: Session) -> dict[str, dict]:
    from app.services import report_studio as RS
    q = RS.QueryIn(
        dataset="mos_capex_summary_calculated",
        dimensions=[
            "__row_order", "__drilldown_key", "sn", "category", "project_count", "total_cost",
            "exp_upto_last_fy", "capex_re", "exp_current_fy", "total_exp", "delay_profile",
        ],
        sort=[RS.SortSpec(by="__row_order", dir="asc")],
        limit=20,
    )
    data = _run_query(db, q)
    return {str(row.get("__drilldown_key")): row for row in data["rows"]}


def _live_mos_rows(db: Session, month: str) -> dict[str, dict]:
    from app.api.v1.mos_reports import mos_capex_summary
    data = mos_capex_summary(report_month=month, db=db)
    by_no = {str(row.get("no") or ""): row for row in data.get("rows", [])}
    by_category = {str(row.get("category") or "").strip().lower(): row for row in data.get("rows", [])}
    total_new = next((row for row in data.get("rows", [])
                      if str(row.get("category") or "").startswith("Total New projects")), None)
    return {
        "1a": by_no.get("1"),
        "1b": by_no.get("2"),
        "1": by_no.get("3"),
        "2": next((row for row in data.get("rows", []) if "milestone payments" in str(row.get("category") or "").lower()), None),
        "3a": by_no.get("3a"),
        "3b": by_no.get("3b"),
        "3": total_new,
        "4": by_category.get("spares & capital repairs"),
        "5": by_category.get("other schemes/ jvs"),
        "total": by_category.get("total"),
    }


def _norm_summary(row: dict | None, source: str) -> dict:
    if not row:
        return {"project_count": None, "total_cost": None, "exp_upto_last_fy": None,
                "capex_re": None, "exp_current_fy": None, "total_exp": None}
    if source == "target":
        return {
            "project_count": row.get("project_count"),
            "total_cost": row.get("total_cost"),
            "exp_upto_last_fy": row.get("exp_upto_last_fy"),
            "capex_re": row.get("capex_re"),
            "exp_current_fy": row.get("exp_current_fy"),
            "total_exp": row.get("total_exp"),
        }
    return {
        "project_count": row.get("projects"),
        "total_cost": row.get("totalCost"),
        "exp_upto_last_fy": row.get("expenditureLastFy"),
        "capex_re": row.get("capexCurrentFy"),
        "exp_current_fy": row.get("expenditureCurrentFy"),
        "total_exp": row.get("totalExpenditure"),
    }


def _diff_value(current, target):
    if current is None and target is None:
        return None
    c = float(current or 0)
    t = float(target or 0)
    return round(c - t, 2)


def _capex_pack_defs() -> list[dict]:
    from datetime import date
    today = date.today()
    fy_year = today.year if today.month >= 4 else today.year - 1
    fy_label = f"{fy_year}-{str(fy_year + 1)[2:]}"
    money = ["gross", "exp_last_fy", "be_fy", "actual_fy", "total_exp"]

    def ms(fields):
        return [{"field": f} for f in fields]

    r1 = {
        "name": f"Physical & Financial Progress of CAPEX Projects — FY {fy_label}",
        "description": "MoS format: quarterly CAPEX overview, category summary, "
                       "project-wise ongoing detail and new projects under consideration. "
                       "All figures in Rs Cr.",
        "category": CAPEX_PACK,
        "sections": [
            {"title": "MoS CAPEX Summary Format - FY 2025-26 up to Mar'26",
             "note": "Matched to the corporate MoS CAPEX format shared in the reference workbook.",
             "spec": {"dataset": "mos_capex_summary_calculated",
                      "dimensions": ["__row_order", "__drilldown_key", "sn", "category", "project_count", "total_cost",
                                     "exp_upto_last_fy", "capex_re", "exp_current_fy", "total_exp",
                                     "delay_profile"],
                      "sort": [{"by": "__row_order", "dir": "asc"}],
                      "limit": 20}},
            {"title": "Overview of CAPEX Progress — BE vs Actual by Month",
             "note": "All figures in Rs. crs. Quarter totals and annual total computed live.",
             "spec": {"dataset": "capex_monthly", "dimensions": ["month_label"],
                      "measures": ms(["be", "actual"]),
                      "pivot": {"on": "month_label", "row_total": True, "quarter_totals": True},
                      "limit": 5000}},
            {"title": "Overview of CAPEX Projects by Category",
             "spec": {"dataset": "pf_projects", "dimensions": ["mos_category"],
                      "measures": ms(["project_count"] + money),
                      "sort": [{"by": "mos_category", "dir": "asc"}],
                      "grand_total": True, "limit": 100}},
            {"title": "Ongoing Projects >= Rs 50 Cr — Physical & Financial Progress",
             "note": "Physical %: i. till last FY / ii. FY plan / iii. FY actual (weighted).",
             "spec": {"dataset": "pf_projects",
                      "dimensions": ["scheme_name", "approval_date", "award_date",
                                     "original_completion", "anticipated_completion", "reason"],
                      "measures": ms(["gross", "phys_last_fy", "phys_fy_plan", "phys_fy_actual",
                                      "exp_last_fy", "be_fy", "actual_fy", "total_exp"]),
                      "filters": {"op": "AND", "conditions": [
                          {"field": "status", "op": "in", "value": ["ongoing", "on_hold"]},
                          {"field": "cost_band", "op": "starts_with", "value": "A"}]},
                      "sort": [{"by": "gross", "dir": "desc"}], "limit": 200}},
            {"title": "Projects < Rs 50 Cr (grouped)",
             "spec": {"dataset": "pf_projects", "dimensions": ["cost_band"],
                      "measures": ms(["project_count"] + money),
                      "filters": {"op": "AND", "conditions": [
                          {"field": "status", "op": "in", "value": ["ongoing", "on_hold"]},
                          {"field": "cost_band", "op": "starts_with", "value": "B"}]},
                      "limit": 10}},
            {"title": "New CAPEX Projects Under Consideration",
             "spec": {"dataset": "pf_projects",
                      "dimensions": ["mos_category", "scheme_name", "approval_date", "award_date"],
                      "measures": ms(["gross", "be_fy", "actual_fy"]),
                      "filters": {"op": "AND", "conditions": [
                          {"field": "mos_category", "op": "starts_with", "value": "3"}]},
                      "sort": [{"by": "mos_category", "dir": "asc"}],
                      "grand_total": True, "limit": 200}},
        ],
    }

    r2 = {
        "name": f"CAPEX Monitoring — Month-wise (FY {fy_label})",
        "description": "Scheme/head rows with BE, RE and Actual spread across the 12 FY months "
                       "plus row totals — the month-wise monitoring format.",
        "category": CAPEX_PACK,
        "sections": [
            {"title": "CAPEX Month-wise — BE / RE / Actual by Row",
             "note": "Rs in crore. Columns: month x (BE, RE, Actual); Total = full-year sum.",
             "spec": {"dataset": "capex_monthly", "dimensions": ["row_name", "month_label"],
                      "measures": ms(["be", "re", "actual"]),
                      "pivot": {"on": "month_label", "row_total": True},
                      "grand_total": True, "limit": 5000}},
            {"title": "Physical Progress Month-wise (Weighted %)",
             "note": "Weighted plan vs actual % per scheme per month, current plans.",
             "spec": {"dataset": "physical_monthly", "dimensions": ["scheme_name", "month_label"],
                      "measures": ms(["plan_pct_w", "actual_pct_w"]),
                      "filters": {"op": "AND", "conditions": [
                          {"field": "fy", "op": "=", "value": fy_year}]},
                      "pivot": {"on": "month_label", "row_total": True}, "limit": 5000}},
        ],
    }

    r3 = {
        "name": f"CAPEX Status of Projects — Sanctioned & New (MoS Backup, FY {fy_label})",
        "description": "MoS presentation backup: per-category and per-scheme sanctioned cost, "
                       "expenditure till last FY, BE/RE, FY expenditure and balance for completion.",
        "category": CAPEX_PACK,
        "sections": [
            {"title": "CAPEX Status by Category & Scheme",
             "note": "balance_for_completion = Total cost - Cumulative expenditure.",
             "spec": {"dataset": "pf_projects", "dimensions": ["mos_category", "scheme_name"],
                      "measures": ms(["gross", "exp_last_fy", "be_fy", "re_fy", "actual_fy", "total_exp"]),
                      "computed": [{"alias": "balance_for_completion",
                                    "expression": "gross - exp_last_fy - actual_fy"}],
                      "sort": [{"by": "mos_category", "dir": "asc"}],
                      "grand_total": True, "limit": 500}},
            {"title": "Quarterly Expenditure — BE / RE / Actual",
             "spec": {"dataset": "capex_monthly", "dimensions": ["quarter"],
                      "measures": ms(["be", "re", "actual"]),
                      "pivot": {"on": "quarter", "row_total": True}, "limit": 100}},
            {"title": "Delay Profile of Ongoing Projects",
             "spec": {"dataset": "pf_projects", "dimensions": ["delay_bucket"],
                      "measures": ms(["project_count", "gross", "total_exp"]),
                      "filters": {"op": "AND", "conditions": [
                          {"field": "status", "op": "in", "value": ["ongoing", "on_hold"]}]},
                      "sort": [{"by": "delay_bucket", "dir": "asc"}],
                      "grand_total": True, "limit": 10}},
        ],
    }
    return [r1, r2, r3]


@router.post("/reports/seed-capex-pack")
def seed_capex_pack(db: Session = Depends(get_db)):
    """(Re)create the 3 standard CAPEX physical-financial reports. Idempotent."""
    _ensure_reports_table(db)
    _ensure_mos_capex_source_rows(db)
    created = []
    for d in _capex_pack_defs():
        payload = ReportIn(**d)
        _validate_sections(payload)
        existing = db.execute(text(
            "SELECT report_id FROM rs_reports WHERE name = :n AND category = :c"),
            {"n": d["name"], "c": CAPEX_PACK}).scalar()
        sections_json = json.dumps([s.model_dump(mode="json") for s in payload.sections])
        if existing:
            db.execute(text(
                "UPDATE rs_reports SET description=:d, sections=CAST(:s AS jsonb), "
                "updated_at=now() WHERE report_id=:r"),
                {"d": d["description"], "s": sections_json, "r": existing})
            created.append({"report_id": existing, "name": d["name"], "updated": True})
        else:
            rid = db.execute(text(
                "INSERT INTO rs_reports (name, description, category, sections) "
                "VALUES (:n, :d, :c, CAST(:s AS jsonb)) RETURNING report_id"),
                {"n": d["name"], "d": d["description"], "c": CAPEX_PACK,
                 "s": sections_json}).scalar()
            created.append({"report_id": rid, "name": d["name"], "updated": False})
    db.commit()
    return {"reports": created}


@router.get("/reports/mos-capex-summary/drilldown/{row_key}")
def mos_capex_summary_drilldown(row_key: str, db: Session = Depends(get_db)):
    _ensure_mos_capex_source_rows(db)
    keys = DRILLDOWN_MAP.get(row_key)
    if not keys:
        raise HTTPException(status_code=404, detail="Drilldown row not found")
    rows = db.execute(text("""
        SELECT source_row, sn, category, project_count, total_cost,
               exp_upto_last_fy, capex_re, exp_current_fy, total_exp,
               source_file, source_sheet
        FROM mos_capex_source_rows
        WHERE report_key = 'fy25_26_mos_capex'
          AND row_key = ANY(:keys)
        ORDER BY source_row
    """), {"keys": keys}).mappings().all()
    out = []
    for r in rows:
        item = {}
        for key, value in dict(r).items():
            if hasattr(value, "isoformat"):
                item[key] = value.isoformat()
            elif value is None or isinstance(value, (int, float, str, bool)):
                item[key] = value
            else:
                item[key] = float(value)
        out.append(item)
    return {
        "row_key": row_key,
        "columns": [
            {"key": "source_row", "label": "Source Row", "type": "int"},
            {"key": "category", "label": "Breakup", "type": "text"},
            {"key": "project_count", "label": "Projects", "type": "number"},
            {"key": "total_cost", "label": "Total Cost", "type": "money"},
            {"key": "exp_upto_last_fy", "label": "Exp up to FY24-25", "type": "money"},
            {"key": "capex_re", "label": "CAPEX FY25-26 RE", "type": "money"},
            {"key": "exp_current_fy", "label": "Exp FY25-26 till Mar'26", "type": "money"},
            {"key": "total_exp", "label": "Total Expenditure", "type": "money"},
        ],
        "rows": out,
    }


@router.get("/reports/mos-capex-summary/reconcile")
def mos_capex_summary_reconcile(
    month: str = "2026-07",
    db: Session = Depends(get_db),
):
    """Compare matched MoS reference figures with current core DB calculation.

    Use this before switching the report from reference-source rows to core live
    tables. Positive variance means current DB is higher than target.
    """
    _ensure_mos_capex_source_rows(db)
    targets = _source_summary_rows(db)
    live = _live_mos_rows(db, month)
    rows = []
    fields = [
        ("project_count", "Projects"),
        ("total_cost", "Total Cost"),
        ("exp_upto_last_fy", "Exp. up to last FY"),
        ("capex_re", "CAPEX RE/BE"),
        ("exp_current_fy", "Exp. current FY"),
        ("total_exp", "Total Expenditure"),
    ]
    for key, (label, _live_key) in MOS_RECON_ROW_MAP.items():
        target = _norm_summary(targets.get(key), "target")
        current = _norm_summary(live.get(key), "live")
        variances = {field: _diff_value(current.get(field), target.get(field)) for field, _ in fields}
        rows.append({
            "row_key": key,
            "label": label,
            "target": target,
            "current": current,
            "variance": variances,
            "matches": all((value is None or abs(value) < 0.01) for value in variances.values()),
        })
    return {
        "month": month,
        "note": "Current = core DB MoS calculation. Target = matched old-report source rows stored in PostgreSQL.",
        "fields": [{"key": key, "label": label} for key, label in fields],
        "rows": rows,
    }


# ================================================================ dashboards
# Power BI-style dashboard canvas: pages of visuals + slicers, persisted as
# query specs (never raw SQL, never frozen numbers). Rendered client-side via
# POST /query/batch so a whole page refreshes in a single round trip.

class DashLayout(BaseModel):
    x: int = 0
    y: int = 0
    w: int = 6
    h: int = 5


class DashVisual(BaseModel):
    id: str
    title: str = ""
    dataset: str
    viz: str = "bar"                 # table|bar|stackedbar|line|area|pie|donut|kpi
    spec: RS.QueryIn
    layout: DashLayout = DashLayout()
    options: dict[str, Any] = {}


class DashSlicer(BaseModel):
    id: str
    dataset: str
    field: str
    label: str = ""
    type: str = "list"               # list | daterange


class DashPage(BaseModel):
    id: str
    title: str = "Page"
    slicers: list[DashSlicer] = []
    visuals: list[DashVisual] = []


class DashboardIn(BaseModel):
    name: str
    description: Optional[str] = None
    pages: list[DashPage] = []
    is_pinned: bool = False


def _ensure_dashboards_table(db: Session) -> None:
    db.execute(text(
        "CREATE TABLE IF NOT EXISTS rs_dashboards ("
        " dashboard_id SERIAL PRIMARY KEY,"
        " name TEXT NOT NULL,"
        " description TEXT,"
        " pages JSONB NOT NULL DEFAULT '[]'::jsonb,"
        " is_pinned BOOLEAN NOT NULL DEFAULT FALSE,"
        " created_by INTEGER,"
        " created_at TIMESTAMPTZ NOT NULL DEFAULT now(),"
        " updated_at TIMESTAMPTZ NOT NULL DEFAULT now())"
    ))
    db.commit()


def _validate_dashboard(payload: DashboardIn) -> None:
    """Every visual spec must compile; every slicer field must exist."""
    for pg in payload.pages:
        for v in pg.visuals:
            if v.spec.dataset != v.dataset:
                raise HTTPException(400, f"Visual '{v.title or v.id}': spec dataset mismatch")
            try:
                RS.compile_query(v.spec)
            except RS.CompileError as e:
                raise HTTPException(400, f"Visual '{v.title or v.id}': {e}")
        for s in pg.slicers:
            try:
                RS.compile_field_values(s.dataset, s.field, None, 1)
            except RS.CompileError as e:
                raise HTTPException(400, f"Slicer '{s.label or s.field}': {e}")


def _dashboard_row(db: Session, dashboard_id: int) -> dict:
    row = db.execute(text(
        "SELECT * FROM rs_dashboards WHERE dashboard_id = :d"), {"d": dashboard_id}
    ).mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Dashboard not found")
    out = dict(row)
    if isinstance(out.get("pages"), str):
        out["pages"] = json.loads(out["pages"])
    return out


@router.get("/dashboards")
def list_dashboards(db: Session = Depends(get_db)):
    _ensure_dashboards_table(db)
    rows = db.execute(text(
        "SELECT dashboard_id, name, description, is_pinned, "
        "       jsonb_array_length(pages) AS page_count, updated_at "
        "FROM rs_dashboards ORDER BY is_pinned DESC, updated_at DESC"
    )).mappings().all()
    return {"dashboards": [dict(r) for r in rows]}


@router.post("/dashboards")
def create_dashboard(payload: DashboardIn, db: Session = Depends(get_db)):
    _ensure_dashboards_table(db)
    _validate_dashboard(payload)
    did = db.execute(text(
        "INSERT INTO rs_dashboards (name, description, pages, is_pinned) "
        "VALUES (:n, :d, CAST(:p AS jsonb), :pin) RETURNING dashboard_id"
    ), {"n": payload.name, "d": payload.description, "pin": payload.is_pinned,
        "p": json.dumps([p.model_dump(mode="json") for p in payload.pages])}).scalar()
    db.commit()
    return {"dashboard_id": did}


@router.get("/dashboards/{dashboard_id}")
def get_dashboard(dashboard_id: int, db: Session = Depends(get_db)):
    _ensure_dashboards_table(db)
    return _dashboard_row(db, dashboard_id)


@router.put("/dashboards/{dashboard_id}")
def update_dashboard(dashboard_id: int, payload: DashboardIn, db: Session = Depends(get_db)):
    _ensure_dashboards_table(db)
    _dashboard_row(db, dashboard_id)
    _validate_dashboard(payload)
    db.execute(text(
        "UPDATE rs_dashboards SET name=:n, description=:d, is_pinned=:pin, "
        "pages=CAST(:p AS jsonb), updated_at=now() WHERE dashboard_id=:r"
    ), {"n": payload.name, "d": payload.description, "pin": payload.is_pinned,
        "p": json.dumps([p.model_dump(mode="json") for p in payload.pages]),
        "r": dashboard_id})
    db.commit()
    return {"ok": True}


@router.delete("/dashboards/{dashboard_id}")
def delete_dashboard(dashboard_id: int, db: Session = Depends(get_db)):
    _ensure_dashboards_table(db)
    _dashboard_row(db, dashboard_id)
    db.execute(text("DELETE FROM rs_dashboards WHERE dashboard_id = :d"), {"d": dashboard_id})
    db.commit()
    return {"ok": True}


@router.post("/dashboards/{dashboard_id}/duplicate")
def duplicate_dashboard(dashboard_id: int, db: Session = Depends(get_db)):
    _ensure_dashboards_table(db)
    src = _dashboard_row(db, dashboard_id)
    did = db.execute(text(
        "INSERT INTO rs_dashboards (name, description, pages) "
        "VALUES (:n, :d, CAST(:p AS jsonb)) RETURNING dashboard_id"
    ), {"n": f"{src['name']} (copy)", "d": src.get("description"),
        "p": json.dumps(src["pages"])}).scalar()
    db.commit()
    return {"dashboard_id": did}


# ---------------------------------------------------------------- batch query

class BatchQueryIn(BaseModel):
    queries: list[RS.QueryIn]


@router.post("/query/batch")
def run_query_batch(payload: BatchQueryIn, db: Session = Depends(get_db)):
    """Run up to 24 structured queries in one call — powers dashboard page
    rendering (all visuals refresh in a single round trip). Per-query errors
    are returned inline so one bad visual never blanks the whole page."""
    if len(payload.queries) > 24:
        raise HTTPException(400, "Batch limited to 24 queries")
    results: list[dict[str, Any]] = []
    for q in payload.queries:
        try:
            results.append({"ok": True, **_run_query(db, q)})
        except HTTPException as e:
            results.append({"ok": False, "error": str(e.detail),
                            "columns": [], "rows": [], "row_count": 0})
    return {"results": results}
