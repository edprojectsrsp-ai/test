"""
Document ingestion pipeline.

For each uploaded document:
  1. Extract text (PDF, DOCX, TXT, images via OCR)
  2. Generate metadata (auto-summary, keywords, important points) via LLM
  3. Chunk into 512-token segments with 50-token overlap
  4. Embed each chunk
  5. Store in document_chunks + document_embeddings

Usage:
    from app.ingestion.processor import process_document
    await process_document(document_id=42)
"""
import os
import json
import hashlib
import logging
from typing import Optional
import psycopg2
import psycopg2.extras
from app.providers.router import get_router
from app.providers.base import ChatMessage
from app.services.embeddings_service import embed_text

logger = logging.getLogger(__name__)


# ============================================================================
# TEXT EXTRACTION
# ============================================================================

def extract_text_from_pdf(file_path: str) -> tuple[str, int, bool]:
    """Extract text from PDF. Returns (text, page_count, needs_ocr)."""
    try:
        import pypdf
        reader = pypdf.PdfReader(file_path)
        page_count = len(reader.pages)
        texts = []
        for page in reader.pages:
            t = page.extract_text() or ""
            texts.append(t)
        full = "\n\n".join(texts).strip()
        # Heuristic: if avg chars per page < 50, probably scanned
        needs_ocr = len(full) < (page_count * 50)
        return full, page_count, needs_ocr
    except Exception as e:
        logger.warning(f"PDF extract failed: {e}")
        return "", 0, True


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from .docx file."""
    try:
        from docx import Document
        doc = Document(file_path)
        paras = [p.text for p in doc.paragraphs]
        # Also tables
        for table in doc.tables:
            for row in table.rows:
                paras.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(p for p in paras if p.strip())
    except Exception as e:
        logger.warning(f"DOCX extract failed: {e}")
        return ""


def ocr_image_or_pdf(file_path: str) -> str:
    """Run OCR via Tesseract. Requires: apt-get install tesseract-ocr poppler-utils + pip pytesseract pdf2image"""
    try:
        from pdf2image import convert_from_path
        import pytesseract
        if file_path.lower().endswith(".pdf"):
            imgs = convert_from_path(file_path, dpi=200)
        else:
            from PIL import Image
            imgs = [Image.open(file_path)]
        texts = [pytesseract.image_to_string(img) for img in imgs]
        return "\n\n".join(texts)
    except Exception as e:
        logger.warning(f"OCR failed: {e}")
        return ""


def extract_text(file_path: str, mime_type: Optional[str] = None) -> tuple[str, int, bool]:
    """Extract text from any supported file. Returns (text, page_count, ocr_was_used)."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        text, pages, needs_ocr = extract_text_from_pdf(file_path)
        if needs_ocr:
            logger.info(f"PDF appears scanned, running OCR: {file_path}")
            ocr_text = ocr_image_or_pdf(file_path)
            if ocr_text:
                return ocr_text, pages, True
        return text, pages, False
    if ext in (".docx",):
        return extract_text_from_docx(file_path), 0, False
    if ext in (".txt", ".md"):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read(), 0, False
    if ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp"):
        return ocr_image_or_pdf(file_path), 1, True
    logger.warning(f"Unsupported file type: {ext}")
    return "", 0, False


# ============================================================================
# CHUNKING
# ============================================================================

def chunk_text(text: str, chunk_size: int = 512, overlap: int = 50) -> list[dict]:
    """Split text into chunks. Returns list of {chunk_no, text, page_number(None for now)}.

    Simple word-based chunking. Good enough for RAG; smarter chunking (sentence/section
    aware) can be added later.
    """
    if not text or not text.strip():
        return []
    words = text.split()
    chunks = []
    chunk_no = 0
    i = 0
    while i < len(words):
        slice_end = min(i + chunk_size, len(words))
        chunk_words = words[i:slice_end]
        chunks.append({
            "chunk_no": chunk_no,
            "text": " ".join(chunk_words),
            "tokens": len(chunk_words),  # rough — words ≠ tokens but good enough
        })
        chunk_no += 1
        if slice_end >= len(words): break
        i = slice_end - overlap
    return chunks


# ============================================================================
# METADATA GENERATION (via LLM)
# ============================================================================

async def generate_metadata(text: str, title: str = "") -> dict:
    """Use Gemini (cheap, big context) to extract summary, keywords, important points."""
    router = get_router()
    snippet = text[:30_000]  # 30K char cap
    prompt = f"""You are processing a project document for an Indian steel plant project management system.

Document title: {title}
Document content (truncated):
---
{snippet}
---

Return a JSON object with these keys (no markdown, no commentary):
- "auto_summary": 2-3 sentences capturing the essence
- "keywords": array of 5-10 relevant tags (lowercase, hyphenated)
- "important_points": array of 3-7 key points as bullet phrases
- "document_date": ISO date if mentioned anywhere in the doc (YYYY-MM-DD), else null
- "parties_mentioned": array of company/agency names mentioned

Output JSON only:
"""
    resp = await router.call(
        [ChatMessage(role="user", content=prompt)],
        task_type="analysis", temperature=0.2, max_tokens=1024,
    )
    if resp.error or not resp.content:
        return {"auto_summary": None, "keywords": [], "important_points": [], "document_date": None}
    try:
        text_out = resp.content.strip()
        # Strip markdown fences if present
        if text_out.startswith("```"):
            text_out = text_out.split("```")[1]
            if text_out.startswith("json"):
                text_out = text_out[4:]
        return json.loads(text_out.strip())
    except Exception as e:
        logger.warning(f"Failed to parse metadata JSON: {e} / content={resp.content[:200]}")
        return {"auto_summary": resp.content[:300], "keywords": [], "important_points": [], "document_date": None}


# ============================================================================
# THE ORCHESTRATOR
# ============================================================================

def get_db():
    dsn = os.environ.get("PROJECT_BRAIN_DB_URL",
                        "postgresql://postgres:abc123@127.0.0.1:5433/project_brain")
    return psycopg2.connect(dsn)


async def process_document(document_id: int) -> dict:
    """Full pipeline for one document. Idempotent — safe to re-run."""
    conn = get_db()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    cur.execute("SELECT * FROM documents WHERE document_id = %s", (document_id,))
    doc = cur.fetchone()
    if not doc:
        conn.close()
        return {"error": f"Document {document_id} not found"}

    file_path = doc["file_path"]
    # If file_path is relative, prepend upload dir
    if not os.path.isabs(file_path):
        upload_dir = os.environ.get("UPLOAD_DIR", "/var/lib/project_brain/uploads")
        file_path = os.path.join(upload_dir, os.path.basename(file_path))
    if not os.path.exists(file_path):
        cur.execute("UPDATE documents SET extraction_status='failed' WHERE document_id=%s", (document_id,))
        conn.commit(); conn.close()
        return {"error": f"File not found: {file_path}"}

    # Mark processing
    cur.execute("UPDATE documents SET extraction_status='processing' WHERE document_id=%s", (document_id,))
    conn.commit()

    try:
        # 1. Extract text
        text, page_count, ocr_used = extract_text(file_path, doc.get("mime_type"))
        if not text.strip():
            cur.execute("UPDATE documents SET extraction_status='failed' WHERE document_id=%s", (document_id,))
            conn.commit(); conn.close()
            return {"error": "No text extracted"}

        # 2. Generate metadata
        meta = await generate_metadata(text, doc["title"])

        # 3. Chunk
        chunks = chunk_text(text, chunk_size=400, overlap=40)

        # 4. Save chunks
        cur.execute("DELETE FROM document_chunks WHERE document_id=%s", (document_id,))
        for c in chunks:
            cur.execute("""
                INSERT INTO document_chunks (document_id, chunk_no, chunk_text, chunk_tokens)
                VALUES (%s, %s, %s, %s)
            """, (document_id, c["chunk_no"], c["text"], c["tokens"]))
        conn.commit()

        # 5. Embed each chunk
        cur.execute("SELECT chunk_id, chunk_text FROM document_chunks WHERE document_id=%s ORDER BY chunk_no",
                    (document_id,))
        chunk_rows = cur.fetchall()
        embeddings_done = 0
        for cr in chunk_rows:
            vec = embed_text(cr["chunk_text"])
            if vec is None: continue
            cur.execute("""
                INSERT INTO document_embeddings (chunk_id, embedding_model, embedding_dim, embedding)
                VALUES (%s, %s, %s, %s::vector)
                ON CONFLICT (chunk_id, embedding_model) DO UPDATE SET embedding=EXCLUDED.embedding
            """, (cr["chunk_id"], os.environ.get("EMBED_MODEL", "all-mpnet-base-v2"),
                  len(vec), str(vec)))
            embeddings_done += 1
        conn.commit()

        # 6. Update document with metadata + final status
        cur.execute("""
            UPDATE documents SET
                page_count = COALESCE(%s, page_count),
                auto_summary = %s, important_points = %s, keywords = %s,
                document_date = %s, ocr_required = %s, ocr_completed = %s,
                chunk_count = %s,
                extraction_status = 'done',
                embedding_status = CASE WHEN %s > 0 THEN 'done' ELSE 'failed' END
            WHERE document_id = %s
        """, (
            page_count or None,
            meta.get("auto_summary"),
            meta.get("important_points") or [],
            meta.get("keywords") or [],
            meta.get("document_date") or None,
            ocr_used, ocr_used,
            len(chunks), embeddings_done, document_id,
        ))
        conn.commit()

        return {
            "document_id": document_id,
            "chunks": len(chunks),
            "embeddings": embeddings_done,
            "ocr_used": ocr_used,
            "page_count": page_count,
            "metadata": meta,
        }

    except Exception as e:
        logger.exception(f"Processing document {document_id} failed")
        cur.execute("UPDATE documents SET extraction_status='failed' WHERE document_id=%s", (document_id,))
        conn.commit()
        return {"error": str(e)}
    finally:
        conn.close()
