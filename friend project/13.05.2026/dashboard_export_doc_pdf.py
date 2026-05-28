import json
import sys


def load_payload(path):
    with open(path, "r", encoding="utf-8-sig") as handle:
        return json.load(handle)


def add_section_lines_pdf(story, styles, title, text):
    from reportlab.platypus import Paragraph, Spacer

    story.append(Paragraph(f"<b>{title}</b>", styles["Heading3"]))
    for line in str(text or "").splitlines():
        line = line.strip()
        if line:
            story.append(Paragraph(line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), styles["BodyText"]))
    story.append(Spacer(1, 10))


def export_pdf(payload, output_path):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Image, SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    current_view_image = payload.get("current_view_image")
    page_size = landscape(A4) if current_view_image else A4
    doc = SimpleDocTemplate(output_path, pagesize=page_size, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=30)
    styles = getSampleStyleSheet()
    story = []

    if current_view_image:
        image = Image(current_view_image)
        max_width = doc.width
        max_height = doc.height
        scale = min(max_width / image.imageWidth, max_height / image.imageHeight)
        image.drawWidth = image.imageWidth * scale
        image.drawHeight = image.imageHeight * scale
        story.append(image)
        doc.build(story)
        return

    story.append(Paragraph(payload.get("title", "Executive Summary Dashboard"), styles["Title"]))
    story.append(Spacer(1, 6))
    story.append(Paragraph(f"Project: {payload.get('project_label', '-')}", styles["Heading2"]))
    story.append(Paragraph(f"Financial Year: {payload.get('fy_label', '-')}", styles["BodyText"]))
    story.append(Paragraph(f"Month: {payload.get('month_label', '-')}", styles["BodyText"]))
    story.append(Paragraph(f"Status: {payload.get('status_text', '-')}", styles["BodyText"]))
    story.append(Spacer(1, 10))

    story.append(Paragraph("<b>Project Identity</b>", styles["Heading3"]))
    for line in payload.get("header_lines", []):
        line = str(line or "").strip()
        if line:
            story.append(Paragraph(line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), styles["BodyText"]))
    story.append(Spacer(1, 10))

    add_section_lines_pdf(story, styles, "Physical Progress Summary", payload.get("physical_text", ""))
    add_section_lines_pdf(story, styles, "Stage Status", payload.get("stage_text", ""))
    add_section_lines_pdf(story, styles, "CAPEX Snapshot", payload.get("capex_text", ""))

    story.append(Paragraph("<b>DPR Insights</b>", styles["Heading3"]))
    for item in payload.get("dpr_summary", []):
        story.append(Paragraph(f"• {str(item)}", styles["BodyText"]))
    story.append(Spacer(1, 10))

    critical_rows = payload.get("critical_rows", [])
    if critical_rows:
        story.append(Paragraph("<b>Critical Path Activities</b>", styles["Heading3"]))
        table_data = [["Activity", "Baseline Start", "Baseline Finish", "Current Status", "Delay"]]
        table_data.extend([[str(col) for col in row] for row in critical_rows])
        table = Table(table_data, repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0b3d91")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        story.append(table)
        story.append(Spacer(1, 10))

    missed_rows = payload.get("missed_rows", [])
    if missed_rows:
        story.append(Paragraph("<b>Missed Baseline Activities</b>", styles["Heading3"]))
        table_data = [["Activity", "Type", "Baseline Date", "Current Status"]]
        table_data.extend([[str(col) for col in row] for row in missed_rows])
        table = Table(table_data, repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#7c2d12")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        story.append(table)

    doc.build(story)


def export_docx(payload, output_path):
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    current_view_image = payload.get("current_view_image")
    if current_view_image:
        section = doc.sections[0]
        section.left_margin = Inches(0.35)
        section.right_margin = Inches(0.35)
        section.top_margin = Inches(0.35)
        section.bottom_margin = Inches(0.35)
        usable_width = section.page_width - section.left_margin - section.right_margin
        doc.add_picture(current_view_image, width=usable_width)
        doc.save(output_path)
        return

    doc.add_heading(payload.get("title", "Executive Summary Dashboard"), 0)
    doc.add_paragraph(f"Project: {payload.get('project_label', '-')}")
    doc.add_paragraph(f"Financial Year: {payload.get('fy_label', '-')}")
    doc.add_paragraph(f"Month: {payload.get('month_label', '-')}")
    doc.add_paragraph(f"Status: {payload.get('status_text', '-')}")

    doc.add_heading("Project Identity", level=1)
    for line in payload.get("header_lines", []):
        if str(line).strip():
            doc.add_paragraph(str(line))

    doc.add_heading("Physical Progress Summary", level=1)
    for line in str(payload.get("physical_text", "")).splitlines():
        if line.strip():
            doc.add_paragraph(line)

    doc.add_heading("Stage Status", level=1)
    for line in str(payload.get("stage_text", "")).splitlines():
        if line.strip():
            doc.add_paragraph(line)

    doc.add_heading("CAPEX Snapshot", level=1)
    for line in str(payload.get("capex_text", "")).splitlines():
        if line.strip():
            doc.add_paragraph(line)

    doc.add_heading("DPR Insights", level=1)
    for item in payload.get("dpr_summary", []):
        doc.add_paragraph(str(item), style="List Bullet")

    critical_rows = payload.get("critical_rows", [])
    if critical_rows:
        doc.add_heading("Critical Path Activities", level=1)
        table = doc.add_table(rows=1, cols=5)
        headers = ["Activity", "Baseline Start", "Baseline Finish", "Current Status", "Delay"]
        for idx, heading in enumerate(headers):
            table.rows[0].cells[idx].text = heading
        for row in critical_rows:
            cells = table.add_row().cells
            for idx, value in enumerate(row[:5]):
                cells[idx].text = str(value)

    missed_rows = payload.get("missed_rows", [])
    if missed_rows:
        doc.add_heading("Missed Baseline Activities", level=1)
        table = doc.add_table(rows=1, cols=4)
        headers = ["Activity", "Type", "Baseline Date", "Current Status"]
        for idx, heading in enumerate(headers):
            table.rows[0].cells[idx].text = heading
        for row in missed_rows:
            cells = table.add_row().cells
            for idx, value in enumerate(row[:4]):
                cells[idx].text = str(value)

    doc.save(output_path)


def main():
    if len(sys.argv) != 4:
        raise SystemExit("Usage: dashboard_export_doc_pdf.py <payload.json> <output-path> <pdf|docx>")
    payload = load_payload(sys.argv[1])
    output_path = sys.argv[2]
    export_format = str(sys.argv[3]).strip().lower()
    if export_format == "pdf":
        export_pdf(payload, output_path)
    elif export_format == "docx":
        export_docx(payload, output_path)
    else:
        raise SystemExit(f"Unsupported export format: {export_format}")
    print(f"Exported {export_format}: {output_path}")


if __name__ == "__main__":
    main()
